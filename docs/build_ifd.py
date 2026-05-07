# -*- coding: utf-8 -*-
"""
기관별 데이터 인터페이스 정의서 (Interface Definition Document, IFD)
LH2 액화수소 인수기지 디지털 트윈 자율안전관리 플랫폼
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_default_font(doc, name="맑은 고딕", size=10):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _force_font(run, name="맑은 고딕"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, level=1, color=(0x1F, 0x3A, 0x5F)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _force_font(run)
        run.font.color.rgb = RGBColor(*color)
    return h


def add_para(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _force_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[0]
        run.text = text
    _force_font(run)
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None, header_bg="1F3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        _force_font(run)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            _force_font(run)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def institution_block(doc, color_hex="2E7D32"):
    """기관별 섹션의 헤더 배경 색상 변경용"""
    pass


# ────────────────────────────────────────────────────────────────────
doc = Document()
set_default_font(doc)
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

# ─────────────── 표지 ───────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("액화수소 인수기지\n디지털 트윈 자율안전관리 플랫폼")
r.bold = True; r.font.size = Pt(22)
_force_font(r)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(20)
r = sub.add_run("기관별 데이터 인터페이스 정의서\n(Interface Definition Document, IFD)")
r.bold = True; r.font.size = Pt(18)
_force_font(r)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(160)
for line in [
    "문서 버전 : v1.0 (PoC 기반 본개발 인터페이스 합의안)",
    "작성 일자 : 2026-05-07",
    "작성 주체 : 유엔이 (플랫폼 개발사)",
    "검토 대상 : KOGAS · KGS · KETI · 세이프티아",
    "동반 문서 : 상세 요구사항 정의서 v1.0 (REQ_SPEC)",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(line); rr.font.size = Pt(11); _force_font(rr)

doc.add_page_break()

# ─────────────── 1. 문서 개요 ───────────────
add_heading(doc, "1. 문서 개요", level=1)

add_heading(doc, "1.1 목적", level=2)
add_para(doc,
    "본 문서는 「액화수소 인수기지 디지털 트윈 자율안전관리 플랫폼」(이하 \"플랫폼\")과 외부기관 "
    "(KOGAS·KGS·KETI·세이프티아) 간의 데이터 교환 인터페이스를 정의한다. PoC 단계의 Mock "
    "Provider 응답 스펙을 기준으로 하되, 본개발 단계에서 각 기관이 자체 서버에서 구현해야 할 "
    "endpoint·요청 파라미터·응답 스키마·비기능 요구사항을 명확히 한다."
)
add_para(doc,
    "각 기관은 본 문서를 검토하여 \"우리 기관에서 제공할 수 있는 데이터·필드·정밀도·응답시간\"을 "
    "확정하고, 필요시 별도 부록으로 추가 요구사항(인증·보안·SLA)을 명시한다."
)

add_heading(doc, "1.2 범위", level=2)
add_para(doc, "(1) 본 문서가 다루는 인터페이스", bold=True)
for t in [
    "플랫폼 → 외부기관 (요청, request)",
    "외부기관 → 플랫폼 (응답, response 또는 비동기 push)",
    "Health Check / 상호 모니터링",
    "에러 처리, 재시도, fallback 정책",
    "데이터 보안 / 인증 / 전송 무결성",
]:
    add_bullet(doc, t)

add_para(doc, "(2) 본 문서가 다루지 않는 범위 (별도 기관 책임)", bold=True)
for t in [
    "외부기관 시스템 내부 구조 (AI 모델 학습, 시뮬레이션 알고리즘 등)",
    "기관 내부 운영 정책",
    "타 시스템과의 인터페이스 (외부기관 ↔ 운영기관 외 시스템)",
]:
    add_bullet(doc, t)

add_heading(doc, "1.3 인터페이스 원칙", level=2)
add_table(doc,
    ["원칙", "내용"],
    [
        ["표준 프로토콜", "REST over HTTPS, JSON 본문 (UTF-8)"],
        ["식별자 통일", "scenario_id (또는 본개발 시 site_id + event_id), equipment_id 표준 코드 공유"],
        ["멱등성", "GET 호출은 멱등, POST 호출은 request_id 기반 중복 방지"],
        ["버전 관리", "URL 경로에 버전 명시 (예: /v1/diagnose) — 본개발 단계 결정"],
        ["시간 표기", "ISO 8601 (YYYY-MM-DDTHH:MM:SS+09:00, KST)"],
        ["문자 인코딩", "UTF-8, 한글 응답 허용"],
        ["단위", "SI 단위 (bar(g), °C, m³/h, kg/h, mm/s, A, %)"],
        ["응답 시간", "기관별 SLA 합의 (본 문서 §4~7 각 기관별 절 참조)"],
        ["에러 응답", "RFC 7807 Problem Details 권장 (type, title, status, detail, instance)"],
        ["인증", "PoC: 미적용 / 본개발: API Key 또는 mTLS (§8 참조)"],
    ],
    col_widths=[3.5, 13.0]
)

doc.add_page_break()

# ─────────────── 2. 시스템 연동 개요 ───────────────
add_heading(doc, "2. 시스템 연동 개요", level=1)

add_heading(doc, "2.1 연동 토폴로지", level=2)
add_code(doc,
    "                           ┌──────────────────────────┐\n"
    "                           │ KOGAS  : 이상탐지/고장인지 AI │\n"
    "                           ├──────────────────────────┤\n"
    "                           │ KGS    : 상호영향 위험예측    │\n"
    " [플랫폼 (유엔이)]  ⇄  HTTPS │ KETI   : 시뮬레이션 엔진     │\n"
    " (Backend on Railway)        │ 세이프티아 : 사고대응/이력 DB │\n"
    "                           └──────────────────────────┘\n"
    "\n"
    "  ※ 각 기관은 자체 서버를 운영하며, 플랫폼은 REST API로만 연동.\n"
    "  ※ PoC에서는 플랫폼 내부의 Mock Provider가 동일 응답 스펙으로 대체.\n"
    "  ※ 본개발에서는 동일 스펙을 유지한 채 endpoint만 외부기관 서버로 교체."
)

add_heading(doc, "2.2 데이터 전송 방식별 분류", level=2)
add_table(doc,
    ["기관", "PoC 방식", "본개발 권고", "사용자 입력 필요"],
    [
        ["KOGAS", "Pull (플랫폼이 scenario_id로 조회)", "Push (실시간 이상감지 시 KOGAS → 플랫폼 webhook) + Pull(상세 조회)", "불필요 (자동 감지)"],
        ["KGS", "Pull / Manual analyze", "이벤트 연계: 자동 (센서값+설비정보 push) / 수동 분석: REST POST", "수동 시 파라미터 입력"],
        ["KETI", "Pull / Manual simulate", "이벤트 연계: 자동 / 수동 시뮬레이션: REST POST + 비동기 콜백", "수동 시 파라미터 입력"],
        ["세이프티아", "Pull (scenario_id 기준)", "Pull + Push (이벤트 종료 시 플랫폼 → 세이프티아 자동 기록)", "불필요"],
    ],
    col_widths=[2.5, 4.0, 6.5, 4.0]
)

add_heading(doc, "2.3 Health Check 규약", level=2)
add_para(doc, "각 기관은 다음의 Health Check endpoint를 제공한다.", bold=False)
add_code(doc,
    "GET /health    →  200 { \"status\": \"ok\",      \"provider\": \"<KOGAS|KGS|KETI|SAFETIA>\" }\n"
    "                  503 { \"status\": \"degraded\", \"provider\": \"<...>\", \"reason\": \"...\" }\n"
    "                  503 { \"status\": \"down\",     \"provider\": \"<...>\" }"
)
add_para(doc,
    "플랫폼은 매 30초 주기로 폴링하여 상단 ApiStatusBar에 🟢/🟡/🔴 상태로 표시한다."
)

doc.add_page_break()

# ─────────────── 3. 공통 데이터 모델 ───────────────
add_heading(doc, "3. 공통 데이터 모델 / 식별자 표준", level=1)

add_heading(doc, "3.1 시나리오 ID (scenario_id)", level=2)
add_para(doc,
    "PoC에서는 SC-01 ~ SC-08의 8종 시나리오 ID로 모든 기관 데이터가 join된다. 본개발에서는 "
    "다음과 같은 변경이 검토된다."
)
add_table(doc,
    ["식별자", "PoC", "본개발 옵션"],
    [
        ["scenario_id", "SC-01 ~ SC-08", "(A) 시나리오 모드 (시연/훈련용) 잔존 / (B) 실 운영 모드에서는 site_id + event_id 조합으로 대체"],
        ["site_id", "—", "본개발 신규: 인수기지 식별자 (예: \"LH2-INCHEON-01\")"],
        ["event_id", "EventLog.event_id (UUID)", "본개발 유지, 외부기관 호출의 correlation key로 활용 권장"],
    ],
    col_widths=[3.0, 4.0, 10.0]
)

add_heading(doc, "3.2 설비 ID (equipment_id)", level=2)
add_para(doc, "12종 핵심설비. 본개발에서는 운영기관 표준 코드로 매핑 또는 확장.", bold=False)
add_table(doc,
    ["equipment_id", "유형", "공간(zone_id)", "설비명"],
    [
        ["TK-101", "STORAGE_TANK", "Z-STO", "LH2 저장탱크 #1"],
        ["TK-102", "STORAGE_TANK", "Z-STO", "LH2 저장탱크 #2"],
        ["BOG-201", "BOG_COMPRESSOR", "Z-BOG", "BOG 압축기 #1"],
        ["PMP-301", "TRANSFER_PUMP", "Z-PUMP", "LH2 이송펌프 #1"],
        ["VAP-401", "VAPORIZER", "Z-VAP", "기화기 #1"],
        ["PIP-501", "MAIN_PIPE", "Z-PIPE", "메인 이송배관 A"],
        ["VAL-601", "VALVE_STATION", "Z-VAL", "벤트스택 #1"],
        ["VAL-602", "VALVE_STATION", "Z-VAL", "벤트스택 #2"],
        ["REL-701", "RELIQUEFIER", "Z-BOG", "BOG 재액화기 #1"],
        ["SHP-001", "LH2_CARRIER", "Z-BERTH", "LH2 운반선 #1"],
        ["ARM-101", "LOADING_ARM", "Z-ARM", "로딩암 #1"],
        ["SWP-001", "SEAWATER_PUMP", "Z-VAP", "해수펌프 #1"],
    ],
    col_widths=[3.0, 4.0, 3.0, 7.0]
)

add_heading(doc, "3.3 센서 ID (sensor_id) — 명명 규칙", level=2)
add_code(doc,
    "{equipment_id}-{TYPE-3letter}-{seq}\n"
    "  ├─ equipment_id : 설비 ID\n"
    "  ├─ TYPE         : PRE(Pressure), TEM(Temperature), FLO(Flow),\n"
    "  │                 VIB(Vibration), CUR(Current), LEV(Level)\n"
    "  └─ seq          : 01, 02, 03 ...\n"
    "\n"
    "예) TK-101-PRE-01, BOG-201-VIB-04, PMP-301-FLO-03, ARM-101-FLO-01"
)

add_heading(doc, "3.4 공통 응답 헤더 (제안)", level=2)
add_code(doc,
    "Content-Type        : application/json; charset=utf-8\n"
    "X-Provider          : KOGAS | KGS | KETI | SAFETIA\n"
    "X-Request-Id        : <UUIDv4, 멱등성 키>\n"
    "X-Response-Time-ms  : <엔진 처리 시간>\n"
    "X-Schema-Version    : v1.0"
)

add_heading(doc, "3.5 공통 에러 응답", level=2)
add_code(doc,
    "{\n"
    "  \"type\":     \"https://platform.example/errors/{kind}\",\n"
    "  \"title\":    \"<short title>\",\n"
    "  \"status\":   400 | 401 | 403 | 404 | 422 | 500 | 503,\n"
    "  \"detail\":   \"<message>\",\n"
    "  \"instance\": \"<request path>\",\n"
    "  \"provider\": \"KOGAS | KGS | KETI | SAFETIA\",\n"
    "  \"timestamp\": \"<ISO8601>\"\n"
    "}"
)

doc.add_page_break()

# ─────────────── 4. KOGAS 인터페이스 ───────────────
add_heading(doc, "4. KOGAS 인터페이스 — 이상탐지 / 고장인지", level=1, color=(0x1B, 0x5E, 0x20))

add_heading(doc, "4.1 책임 영역", level=2)
for t in [
    "센서 데이터(시계열) 기반 이상감지 알고리즘 운영",
    "고장 종류·고장 부위 분류 (fault_code, fault_name, suspected_part)",
    "진단 신뢰도(diagnosis_confidence) 산출",
    "근거 센서 ID 목록 (sensor_evidence) 제공",
    "(본개발) 실시간 push 또는 학습 모델 재학습용 데이터 수신 정책 결정",
]:
    add_bullet(doc, t)

add_heading(doc, "4.2 Endpoint", level=2)
add_table(doc,
    ["Method", "Path", "용도", "PoC", "본개발"],
    [
        ["GET", "/health", "Health check", "✅", "✅"],
        ["GET", "/diagnose/{scenario_id}", "(시연용) 시나리오별 사전 진단 결과", "✅", "옵션 (시연 모드 잔존)"],
        ["POST", "/diagnose", "실시간 진단 요청 (센서값 push)", "—", "✅ 신규"],
        ["POST (callback)", "{platform}/api/provider/kogas/push", "이상감지 시 KOGAS → 플랫폼 webhook", "—", "✅ 신규 (옵션)"],
    ],
    col_widths=[2.5, 5.5, 5.0, 1.5, 2.5]
)

add_heading(doc, "4.3 [본개발] POST /diagnose — 요청", level=2)
add_code(doc,
    "POST /v1/diagnose HTTP/1.1\n"
    "Content-Type: application/json\n"
    "X-Request-Id: <UUID>\n"
    "Authorization: Bearer <token>\n"
    "\n"
    "{\n"
    "  \"site_id\":     \"LH2-INCHEON-01\",\n"
    "  \"event_id\":    \"<UUID, 선택>\",\n"
    "  \"scenario_id\": \"SC-01\",                 // 시연 모드 시\n"
    "  \"target_equipment_id\": \"BOG-201\",\n"
    "  \"window_sec\":  300,                      // 분석 시간창\n"
    "  \"sensor_data\": [\n"
    "    {\n"
    "      \"sensor_id\":   \"BOG-201-PRE-01\",\n"
    "      \"sensor_type\": \"PRESSURE\",\n"
    "      \"unit\":        \"bar(g)\",\n"
    "      \"samples\": [\n"
    "        { \"timestamp\": \"2026-05-07T10:00:00+09:00\", \"value\": 10.05, \"quality\": \"GOOD\" },\n"
    "        { \"timestamp\": \"2026-05-07T10:00:05+09:00\", \"value\": 10.13, \"quality\": \"GOOD\" }\n"
    "      ]\n"
    "    }\n"
    "    // ... 다중 센서\n"
    "  ]\n"
    "}"
)

add_heading(doc, "4.4 응답 스키마 (PoC Mock과 동일 + 본개발 확장)", level=2)
add_table(doc,
    ["필드", "타입", "필수", "설명", "예시"],
    [
        ["request_id", "string", "✅", "내부 요청 ID (멱등성)", "KOGAS-SC-01"],
        ["scenario_id", "string", "옵션", "시연 모드 식별자", "SC-01"],
        ["event_id", "string (UUID)", "옵션", "본개발: 이벤트 correlation", "—"],
        ["target_equipment_id", "string", "✅", "진단 대상 설비", "BOG-201"],
        ["fault_code", "string", "✅", "표준 고장코드 (KOGAS 정의)", "FLT-01, FLT-ARM-ESD"],
        ["fault_name", "string", "✅", "고장명 (한글)", "BOG 압축기 트립 또는 전동기 과전류 보호동작"],
        ["diagnosis_confidence", "number", "✅", "신뢰도 0~1 (소수점 2자리)", "0.75"],
        ["suspected_part", "string", "✅", "의심 부위 (한글)", "전동기 / 베어링 / 밸브구동부"],
        ["sensor_evidence", "array", "✅", "근거 센서 (string 배열 또는 객체 배열)", "[\"BOG-201-PRE-01\", ...] 또는 [{sensor_id, desc}]"],
        ["timestamp", "ISO8601", "✅(본개발)", "진단 시각", "2026-05-07T10:05:23+09:00"],
        ["model_version", "string", "옵션", "AI 모델 버전", "kogas-v1.2"],
    ],
    col_widths=[3.0, 2.0, 1.0, 5.5, 5.5]
)

add_heading(doc, "4.5 응답 예시 (PoC SC-01)", level=2)
add_code(doc,
    "{\n"
    "  \"request_id\": \"KOGAS-SC-01\",\n"
    "  \"scenario_id\": \"SC-01\",\n"
    "  \"target_equipment_id\": \"BOG-201\",\n"
    "  \"fault_code\": \"FLT-01\",\n"
    "  \"fault_name\": \"BOG 압축기 트립 또는 전동기 과전류 보호동작\",\n"
    "  \"diagnosis_confidence\": 0.75,\n"
    "  \"suspected_part\": \"전동기\",\n"
    "  \"sensor_evidence\": [\n"
    "    \"BOG-201-PRE-01\",\n"
    "    \"BOG-201-TEM-02\",\n"
    "    \"BOG-201-FLO-03\"\n"
    "  ]\n"
    "}"
)
add_para(doc, "응답 예시 (PoC SC-08 — sensor_evidence 객체 변형)", bold=True)
add_code(doc,
    "{\n"
    "  \"request_id\": \"KOGAS-SC-08\",\n"
    "  \"scenario_id\": \"SC-08\",\n"
    "  \"target_equipment_id\": \"ARM-101\",\n"
    "  \"fault_code\": \"FLT-ARM-ESD\",\n"
    "  \"fault_name\": \"로딩암 긴급차단(ESD) 작동, 플랜지 기밀 이상 또는 계류 이탈 감지\",\n"
    "  \"diagnosis_confidence\": 0.82,\n"
    "  \"suspected_part\": \"ESD 밸브, 플랜지 실링, 계류 센서\",\n"
    "  \"sensor_evidence\": [\n"
    "    { \"sensor_id\": \"ARM-101-FLO-01\", \"desc\": \"유량 급감 (120→0 m³/h)\" },\n"
    "    { \"sensor_id\": \"ARM-101-VIB-01\", \"desc\": \"진동 급증 후 정지\" },\n"
    "    { \"sensor_id\": \"ARM-101-PRE-01\", \"desc\": \"압력 급변동\" }\n"
    "  ]\n"
    "}"
)

add_heading(doc, "4.6 KOGAS 합의 필요 항목 (본개발 진입 전)", level=2)
add_table(doc,
    ["#", "항목", "옵션 / 협의 사항"],
    [
        ["1", "fault_code 표준 체계", "현재: FLT-01 ~ FLT-07, FLT-ARM-ESD / 본개발: KOGAS 자체 코드 체계 통일 필요 (예: FLT-{설비유형}-{seq})"],
        ["2", "sensor_evidence 형식", "string 배열 vs 객체 배열 — 본개발에서는 객체 배열 (sensor_id + desc) 통일 권고"],
        ["3", "진단 응답 시간 SLA", "동기 호출: 5초 이내 / 비동기 콜백 허용 여부"],
        ["4", "Push 방식 도입", "이상감지 시 KOGAS → 플랫폼 webhook 적용 여부, retry 정책"],
        ["5", "센서 데이터 push 주기", "전수 센서 vs 이상 의심 센서만 / 5초 vs 1초 vs 100ms"],
        ["6", "재학습 데이터 회신", "플랫폼이 정/오진 라벨링하여 KOGAS로 회신 (모델 재학습용) 여부"],
        ["7", "다중 동시 진단", "동일 설비 다중 동시 진단 시 처리 (큐잉 / 거부 / 병합)"],
        ["8", "신뢰도 산출 기준", "0~1 소수점 2자리 / 분류 분포 (top-1, top-3)"],
    ],
    col_widths=[0.7, 5.0, 11.0]
)

doc.add_page_break()

# ─────────────── 5. KGS 인터페이스 ───────────────
add_heading(doc, "5. KGS 인터페이스 — 상호영향 위험예측", level=1, color=(0x4A, 0x14, 0x8C))

add_heading(doc, "5.1 책임 영역", level=2)
for t in [
    "단일 설비 이상이 다른 설비로 전파되는 영향(impact) 분석",
    "영향 점수(impact_score), 위험 등급(risk_level), 예측 시점(predicted_after_sec) 산출",
    "영향 유형(impact_type) 분류",
    "HAZOP 연계 (hazop_id)",
    "권고조치(recommended_action) 텍스트 제공",
    "2D 네트워크·3D 컬러링용 시각화 hint (color_2d, color_3d)",
]:
    add_bullet(doc, t)

add_heading(doc, "5.2 Endpoint", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/health", "Health check"],
        ["GET", "/analyze/{scenario_id}", "(시연용) 시나리오별 사전 분석 결과 배열"],
        ["POST", "/analyze", "[본개발 핵심] 실시간 위험예측 (센서값+설비정보 입력 → 영향 분석 배열 응답)"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.3 [본개발] POST /analyze — 요청", level=2)
add_code(doc,
    "POST /v1/analyze HTTP/1.1\n"
    "Content-Type: application/json\n"
    "\n"
    "{\n"
    "  \"site_id\":     \"LH2-INCHEON-01\",\n"
    "  \"event_id\":    \"<UUID>\",\n"
    "  \"scenario_id\": \"SC-01\",                  // 시연 모드 시\n"
    "  \"trigger_equipment_id\": \"BOG-201\",\n"
    "  \"sensor_data\": [ ... 동일 포맷 ... ],\n"
    "  \"params\": {\n"
    "    \"analysis_window_min\": 60,             // 0~60분\n"
    "    \"pressure_excess_pct\": 15,\n"
    "    \"temperature_deviation_c\": 5,\n"
    "    \"propagation_depth\": 3                 // 영향 그래프 깊이\n"
    "  }\n"
    "}"
)

add_heading(doc, "5.4 응답 스키마 (배열)", level=2)
add_table(doc,
    ["필드", "타입", "필수", "설명", "예시"],
    [
        ["analysis_id", "string", "✅", "분석 단건 ID", "KGS-SC-01-01"],
        ["scenario_id / event_id", "string", "✅", "correlation key", "SC-01"],
        ["trigger_equipment_id", "string", "✅", "이상 발생 설비", "BOG-201"],
        ["affected_equipment_id", "string", "✅", "영향 받는 설비 (자기 자신 포함 가능)", "TK-101"],
        ["impact_type", "string (enum)", "✅", "영향 유형 (아래 enum)", "PRESSURE_PROPAGATION"],
        ["impact_score", "integer (0~100)", "✅", "영향 점수", "76"],
        ["risk_level", "string (enum)", "✅", "위험 등급", "WARNING"],
        ["predicted_after_sec", "integer", "✅", "예측 시점 (sec, 0이면 즉시)", "180"],
        ["color_2d", "string (hex)", "옵션", "2D 네트워크 색상 hint", "#FFEE58"],
        ["color_3d", "string (enum)", "옵션", "3D 컬러링 hint", "affected"],
        ["hazop_id", "string", "옵션", "연계 HAZOP", "HZ-LH2-001"],
        ["recommended_action", "string", "✅", "권고조치 (한글)", "압축기 정지 확인 후 저장탱크 압력상승률 감시"],
    ],
    col_widths=[3.0, 2.5, 1.0, 5.5, 5.0]
)

add_heading(doc, "5.5 enum 정의", level=2)
add_table(doc,
    ["enum", "허용값"],
    [
        ["impact_type", "PRIMARY_EVENT, PRESSURE_PROPAGATION, PSV_ACTIVATION_RISK, BACKUP_ACTIVATION, FLOW_INTERRUPTION, PRESSURE_BUILDUP, SUPPLY_SHORTAGE, FLOW_REDUCTION (확장 가능)"],
        ["risk_level", "INFO, WARNING, CRITICAL, EMERGENCY"],
        ["color_2d", "#EF5350(critical) / #FFEE58(affected) / #FFA726(warning) / #FF1744(emergency)"],
        ["color_3d", "critical / affected / warning / emergency / simTarget"],
    ],
    col_widths=[3.5, 13.5]
)

add_heading(doc, "5.6 응답 예시 (PoC SC-01 영향 체인)", level=2)
add_code(doc,
    "[\n"
    "  {\n"
    "    \"analysis_id\": \"KGS-SC-01-01\",\n"
    "    \"scenario_id\": \"SC-01\",\n"
    "    \"trigger_equipment_id\":  \"BOG-201\",\n"
    "    \"affected_equipment_id\": \"BOG-201\",\n"
    "    \"impact_type\": \"PRIMARY_EVENT\",\n"
    "    \"impact_score\": 88,\n"
    "    \"risk_level\":   \"CRITICAL\",\n"
    "    \"predicted_after_sec\": 0,\n"
    "    \"color_2d\": \"#EF5350\", \"color_3d\": \"critical\",\n"
    "    \"hazop_id\": \"HZ-LH2-001\",\n"
    "    \"recommended_action\": \"압축기 정지 확인 후 저장탱크 압력상승률 감시\"\n"
    "  },\n"
    "  {\n"
    "    \"analysis_id\": \"KGS-SC-01-02\",\n"
    "    \"trigger_equipment_id\":  \"BOG-201\",\n"
    "    \"affected_equipment_id\": \"TK-101\",\n"
    "    \"impact_type\": \"PRESSURE_PROPAGATION\",\n"
    "    \"impact_score\": 76, \"risk_level\": \"WARNING\",\n"
    "    \"predicted_after_sec\": 180,\n"
    "    \"color_2d\": \"#FFEE58\", \"color_3d\": \"affected\",\n"
    "    \"hazop_id\": \"HZ-LH2-001\",\n"
    "    \"recommended_action\": \"압축기 정지 확인 후 저장탱크 압력상승률 감시\"\n"
    "  },\n"
    "  {\n"
    "    \"analysis_id\": \"KGS-SC-01-03\",\n"
    "    \"affected_equipment_id\": \"REL-701\",\n"
    "    \"impact_type\": \"PRESSURE_PROPAGATION\",\n"
    "    \"impact_score\": 64, \"risk_level\": \"WARNING\",\n"
    "    \"predicted_after_sec\": 360, ...\n"
    "  },\n"
    "  {\n"
    "    \"analysis_id\": \"KGS-SC-01-04\",\n"
    "    \"affected_equipment_id\": \"VAL-601\",\n"
    "    \"impact_type\": \"PSV_ACTIVATION_RISK\",\n"
    "    \"impact_score\": 58, \"risk_level\": \"WARNING\",\n"
    "    \"predicted_after_sec\": 540, ...\n"
    "  }\n"
    "]"
)

add_heading(doc, "5.7 KGS 합의 필요 항목", level=2)
add_table(doc,
    ["#", "항목", "옵션 / 협의 사항"],
    [
        ["1", "impact_type 표준 체계", "현재 8종 / KGS가 추가 정의 가능 (위험 분석 기법에 따라)"],
        ["2", "impact_score 산출 기준", "정량적 기준 (확률·강도·노출시간 기반) 합의 필요"],
        ["3", "predicted_after_sec 정밀도", "초 단위 / 분 단위 / 시간창(min, max) 표현"],
        ["4", "분석 응답 시간 SLA", "동기 5초 이내 (PoC) / 본개발: 영향 그래프 깊이에 따라 SLA 차등"],
        ["5", "영향 그래프 깊이", "PoC: 4단계 / 본개발: 운영기관 P&ID 복잡도에 따라 결정"],
        ["6", "recommended_action 텍스트", "자유 텍스트 vs SOP 매핑 ID 회신 권고"],
        ["7", "color_2d / color_3d 위치", "KGS 응답에 포함 vs 플랫폼이 risk_level 기반 자체 매핑"],
        ["8", "센서 데이터 입력 범위", "전수 센서 / 이상 의심 센서만 / 임계치 초과 센서"],
    ],
    col_widths=[0.7, 5.0, 11.0]
)

doc.add_page_break()

# ─────────────── 6. KETI 인터페이스 ───────────────
add_heading(doc, "6. KETI 인터페이스 — 시뮬레이션 / 의사결정지원", level=1, color=(0xC6, 0x28, 0x28))

add_heading(doc, "6.1 책임 영역", level=2)
for t in [
    "이상 시나리오 시뮬레이션 (압력·온도·유량 등 시간경과 예측)",
    "대응안 추천 (Option A / Option B 또는 N개)",
    "각 대응안별 안정화 시간(stabilization_min), 위험도, 세부 설명",
    "하이브리드 모델링 (물리 + AI)",
    "(본개발 옵션) 동적 저작도구 — 새 시나리오 정의·실행",
]:
    add_bullet(doc, t)

add_heading(doc, "6.2 Endpoint", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/health", "Health check"],
        ["GET", "/simulate/{scenario_id}", "(시연용) 시나리오별 사전 시뮬레이션 결과"],
        ["POST", "/simulate", "[본개발 핵심] 실시간 시뮬레이션 (파라미터 입력 → 결과 응답)"],
        ["POST (callback)", "{platform}/api/provider/keti/result", "비동기 모드: 시뮬레이션 완료 시 KETI → 플랫폼 webhook"],
    ],
    col_widths=[2.5, 5.5, 9.0]
)

add_heading(doc, "6.3 [본개발] POST /simulate — 요청", level=2)
add_code(doc,
    "POST /v1/simulate HTTP/1.1\n"
    "Content-Type: application/json\n"
    "\n"
    "{\n"
    "  \"site_id\":     \"LH2-INCHEON-01\",\n"
    "  \"event_id\":    \"<UUID>\",\n"
    "  \"scenario_id\": \"SC-01\",                  // 시연 모드 시\n"
    "  \"trigger_equipment_id\": \"BOG-201\",\n"
    "  \"params\": {\n"
    "    \"fault_type\":             \"OVERPRESSURE\",\n"
    "    \"initial_pressure_bar\":   12.5,\n"
    "    \"duration_hr\":            2,\n"
    "    \"temperature_deviation_c\": 5,\n"
    "    \"flow_change_pct\":        -40,\n"
    "    \"option_count\":           2              // 1~N 대응안 개수\n"
    "  },\n"
    "  \"return_mode\": \"sync\"                    // sync | async\n"
    "}"
)

add_heading(doc, "6.4 응답 스키마", level=2)
add_table(doc,
    ["필드", "타입", "필수", "설명", "예시"],
    [
        ["simulation_id", "string", "✅", "시뮬레이션 ID", "KETI-SC-01"],
        ["scenario_id / event_id", "string", "✅", "correlation key", "SC-01"],
        ["trigger_equipment_id", "string", "✅", "트리거 설비", "BOG-201"],
        ["simulation_summary", "string", "✅", "사황 요약 (한글, 1~3문장)", "BOG 압축기 트립으로 저장탱크 압력 지속 상승 중. PSV 개방까지 약 18분 소요 예상."],
        ["recommended_option_a", "string", "✅", "1순위 대응안 (한글)", "재액화기 긴급 가동, 압축기 수리 병행"],
        ["recommended_option_b", "string", "옵션", "2순위 대응안 (한글)", "벤트스택 수동 개방, 압축기 정비"],
        ["option_a_stabilization_min", "integer", "✅", "Option A 안정화 시간(분)", "18"],
        ["option_b_stabilization_min", "integer", "옵션", "Option B 안정화 시간(분)", "25"],
        ["option_a_risk", "enum", "✅", "Option A 위험도", "MEDIUM"],
        ["option_b_risk", "enum", "옵션", "Option B 위험도", "LOW"],
        ["option_a_detail", "string", "✅", "Option A 세부 설명", "재액화기 처리용량 한계 / 탱크 압력 11.5bar 이하 안정화"],
        ["option_b_detail", "string", "옵션", "Option B 세부 설명", "BOG 대기 방출 / 방출량 약 500kg/h"],
        ["expected_stabilization_min", "integer", "✅", "권고 대응안 안정화 시간", "18"],
        ["timeseries_forecast", "array (옵션)", "—", "본개발: 시간별 예측 시계열", "[{t, P, T, Flow, ...}, ...]"],
        ["simulation_timestamp", "ISO8601", "✅(본개발)", "시뮬레이션 시각", "—"],
    ],
    col_widths=[3.5, 2.5, 1.0, 5.5, 5.0]
)

add_heading(doc, "6.5 enum 정의", level=2)
add_table(doc,
    ["enum", "허용값"],
    [
        ["option_*_risk", "NONE, LOW, MEDIUM, HIGH"],
        ["fault_type (요청)", "OVERPRESSURE, CAVITATION, STUCK_VALVE, COLD_BRITTLE, BOG_SURGE, ESD_TRIP, NORMAL_OPS (확장 가능)"],
        ["return_mode", "sync (동기 응답) / async (즉시 ack 후 콜백)"],
    ],
    col_widths=[3.5, 13.5]
)

add_heading(doc, "6.6 응답 예시 (PoC SC-01)", level=2)
add_code(doc,
    "{\n"
    "  \"simulation_id\": \"KETI-SC-01\",\n"
    "  \"scenario_id\":   \"SC-01\",\n"
    "  \"trigger_equipment_id\": \"BOG-201\",\n"
    "  \"simulation_summary\": \"BOG 압축기 트립으로 저장탱크 압력이 지속 상승 중. PSV 개방 압력(13.5bar)까지 약 18분 소요 예상. 즉각적인 BOG 처리 경로 확보가 필요함.\",\n"
    "  \"recommended_option_a\": \"재액화기(REL-701) 긴급 가동하여 BOG를 LH2로 재전환. 압축기 수리 병행.\",\n"
    "  \"recommended_option_b\": \"벤트스택(VAL-601) 수동 개방하여 BOG를 대기 방출 후, 압축기 정비 실시.\",\n"
    "  \"option_a_stabilization_min\": 18,\n"
    "  \"option_b_stabilization_min\": 25,\n"
    "  \"option_a_risk\": \"MEDIUM\",\n"
    "  \"option_b_risk\": \"LOW\",\n"
    "  \"option_a_detail\": \"재액화기 처리용량 한계로 잔여 BOG 일부 축적 가능. 탱크 압력 11.5bar 이하로 안정화.\",\n"
    "  \"option_b_detail\": \"BOG 대기 방출로 환경 영향 발생. 방출량 약 500kg/h. 안전밸브 개방 위험 해소.\",\n"
    "  \"expected_stabilization_min\": 18\n"
    "}"
)

add_heading(doc, "6.7 KETI 합의 필요 항목", level=2)
add_table(doc,
    ["#", "항목", "옵션 / 협의 사항"],
    [
        ["1", "대응안 개수", "현재 A/B 고정 / 본개발: 1~N 가변 (option_count 파라미터)"],
        ["2", "응답 시간 SLA", "동기: 30초 이내 권장 / 30초 초과 시 async 모드 필수"],
        ["3", "비동기 콜백 URL", "플랫폼이 KETI에 콜백 URL 등록 (X-Callback-Url 헤더)"],
        ["4", "timeseries_forecast 제공 여부", "PoC: 미제공 / 본개발: 시간별 예측 시계열 (압력/온도/유량) 제공 시 3D 시각화 강화 가능"],
        ["5", "simulation_summary 길이 제한", "최대 N자 (UI 표시 영역 제약)"],
        ["6", "동적 저작도구 연동", "신규 시나리오 정의 → KETI 시뮬레이션 → 플랫폼 등록 워크플로우 합의"],
        ["7", "시뮬레이션 입력 표준", "fault_type enum 통일, 파라미터 단위 표준화"],
        ["8", "결과 영구 저장", "플랫폼 캐시 vs KETI 자체 저장 vs 양쪽 저장"],
    ],
    col_widths=[0.7, 5.0, 11.0]
)

doc.add_page_break()

# ─────────────── 7. 세이프티아 인터페이스 ───────────────
add_heading(doc, "7. 세이프티아 인터페이스 — 사고대응 / 이력관리", level=1, color=(0xE6, 0x5A, 0x00))

add_heading(doc, "7.1 책임 영역", level=2)
for t in [
    "설비별 정비/점검/사고/교체 이력 보관 및 조회",
    "비상상황 시나리오 라이브러리 (linked_sop_id 매핑)",
    "운영자 메모(operator_note) 관리",
    "(본개발) 이벤트 종료 시 플랫폼 → 세이프티아 자동 기록",
    "(본개발) 사고대응 절차 라이브러리 동기화",
]:
    add_bullet(doc, t)

add_heading(doc, "7.2 Endpoint", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/health", "Health check"],
        ["GET", "/history/{scenario_id}", "(시연용) 시나리오별 이력 배열"],
        ["GET", "/history", "[본개발] 이력 검색 (?equipment_id, from, to, type, page, size)"],
        ["GET", "/history/{history_id}", "[본개발] 단건 이력 상세"],
        ["POST", "/history", "[본개발] 이벤트 종료 시 플랫폼 → 세이프티아 신규 이력 등록"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "7.3 [본개발] GET /history — 요청", level=2)
add_code(doc,
    "GET /v1/history?equipment_id=BOG-201&from=2025-01-01&to=2026-12-31\n"
    "    &type=MAINTENANCE&page=0&size=20\n"
    "Authorization: Bearer <token>\n"
    "\n"
    "쿼리 파라미터:\n"
    "  equipment_id : 설비 ID (필수)\n"
    "  from / to    : ISO 날짜 (옵션, 기본 최근 1년)\n"
    "  type         : MAINTENANCE | INSPECTION | INCIDENT | REPLACE | NORMAL_PROFILE (옵션)\n"
    "  page / size  : 페이지네이션 (옵션)\n"
    "  scenario_id  : (시연 모드)"
)

add_heading(doc, "7.4 응답 스키마 (배열 또는 페이지)", level=2)
add_table(doc,
    ["필드", "타입", "필수", "설명", "예시"],
    [
        ["history_id", "string", "✅", "이력 단건 ID", "SAFE-SC-01"],
        ["scenario_id", "string", "옵션", "시연 모드", "SC-01"],
        ["event_id", "string (UUID)", "옵션", "본개발: 연계 이벤트", "—"],
        ["equipment_id", "string", "✅", "대상 설비", "BOG-201"],
        ["history_type", "enum", "✅(본개발)", "이력 유형", "MAINTENANCE"],
        ["last_maintenance_date", "ISO date", "✅", "최근 정비일", "2026-01-18"],
        ["past_incident_summary", "string", "✅", "과거 사고/상태 요약 (한글)", "저장탱크에서 발생한 BOG를 처리하지 못함"],
        ["linked_sop_id", "string", "옵션", "연계 SOP", "SOP-BOG-TRIP-01"],
        ["operator_note", "string", "옵션", "운영자 메모/권고 (한글)", "압축기 진동·전류 실시간 감시, 예비 압축기 준비 ..."],
        ["created_at", "ISO8601", "✅(본개발)", "이력 등록 시각", "—"],
        ["updated_at", "ISO8601", "옵션", "최종 수정 시각", "—"],
    ],
    col_widths=[3.5, 2.5, 1.0, 5.0, 5.0]
)

add_heading(doc, "7.5 [본개발] POST /history — 신규 이력 등록 (플랫폼 → 세이프티아)", level=2)
add_code(doc,
    "POST /v1/history HTTP/1.1\n"
    "Content-Type: application/json\n"
    "\n"
    "{\n"
    "  \"site_id\":      \"LH2-INCHEON-01\",\n"
    "  \"event_id\":     \"<UUID>\",\n"
    "  \"equipment_id\": \"BOG-201\",\n"
    "  \"history_type\": \"INCIDENT\",\n"
    "  \"summary\":      \"BOG 압축기 트립 발생, 18분 내 재액화기 가동으로 안정화\",\n"
    "  \"linked_sop_id\":\"SOP-BOG-TRIP-01\",\n"
    "  \"executed_steps\": [1, 2, 3, 4, 5],\n"
    "  \"operator_note\": \"...\",\n"
    "  \"closed_at\":   \"2026-05-07T11:30:00+09:00\"\n"
    "}\n"
    "\n"
    "응답:\n"
    "  201 Created\n"
    "  { \"history_id\": \"SAFE-2026-05-07-0001\", \"created_at\": \"...\" }"
)

add_heading(doc, "7.6 응답 예시 (PoC SC-01)", level=2)
add_code(doc,
    "{\n"
    "  \"history_id\":  \"SAFE-SC-01\",\n"
    "  \"scenario_id\": \"SC-01\",\n"
    "  \"equipment_id\":\"BOG-201\",\n"
    "  \"last_maintenance_date\": \"2026-01-18\",\n"
    "  \"past_incident_summary\":  \"저장탱크에서 발생한 BOG를 처리하지 못함\",\n"
    "  \"linked_sop_id\": \"SOP-BOG-TRIP-01\",\n"
    "  \"operator_note\": \"압축기 진동·전류 실시간 감시, 예비 압축기/우회라인 준비, 토출측 차압 모니터링 강화\"\n"
    "}"
)

add_heading(doc, "7.7 세이프티아 합의 필요 항목", level=2)
add_table(doc,
    ["#", "항목", "옵션 / 협의 사항"],
    [
        ["1", "history_type 표준 체계", "MAINTENANCE / INSPECTION / INCIDENT / REPLACE / NORMAL_PROFILE 외 추가 구분"],
        ["2", "페이지네이션 / 검색 옵션", "필수 필터 (설비, 기간, 유형) + 정렬 옵션"],
        ["3", "이력 영구 보관 vs 캐시", "(A) 매 호출 / (B) 플랫폼 캐시(TTL) / (C) 부분 동기화"],
        ["4", "이벤트 종료 시 자동 기록", "플랫폼 → 세이프티아 POST /history 적용 여부 + 필수 필드"],
        ["5", "사고대응 라이브러리 연동", "세이프티아 보유 사고대응 절차를 SOP 카탈로그로 가져올지"],
        ["6", "첨부파일", "사진·문서·서명 첨부 시 별도 endpoint"],
        ["7", "응답 시간 SLA", "동기 3초 이내 권장"],
    ],
    col_widths=[0.7, 5.0, 11.0]
)

doc.add_page_break()

# ─────────────── 8. 보안 / 인증 / 운영 ───────────────
add_heading(doc, "8. 보안 / 인증 / 운영 정책", level=1)

add_heading(doc, "8.1 인증 방식 (본개발)", level=2)
add_table(doc,
    ["방식", "설명", "권고"],
    [
        ["API Key", "정적 토큰 헤더(X-API-Key) + IP 화이트리스트", "초기 단계 (Phase 3)"],
        ["OAuth 2.0 (Client Credentials)", "기관별 client_id/secret, JWT 발급", "권고 (운영)"],
        ["Mutual TLS (mTLS)", "양방향 인증서 검증", "고위험 데이터 (기관별 정책)"],
    ],
    col_widths=[5.0, 8.0, 4.0]
)

add_heading(doc, "8.2 전송 보안", level=2)
for t in [
    "TLS 1.2 이상 의무 (1.0/1.1 차단)",
    "암호 스위트 : ECDHE-RSA / ECDHE-ECDSA 계열",
    "인증서 : Let's Encrypt 또는 운영기관 사설 CA",
    "민감 정보(개인정보·서명·IP) 본문 평문 전송 금지 — 필요 시 추가 암호화",
]:
    add_bullet(doc, t)

add_heading(doc, "8.3 재시도 / Backoff / Rate Limit", level=2)
add_table(doc,
    ["항목", "정책"],
    [
        ["재시도 (플랫폼 → 기관)", "Exponential backoff: 1s, 2s, 4s, 8s (최대 4회)"],
        ["타임아웃", "동기 호출 30초, 비동기 ack 5초, Health check 3초"],
        ["Rate Limit (기관 → 플랫폼)", "기관별 합의 (예: 100 req/min/IP)"],
        ["서버 다운 시 fallback", "플랫폼은 마지막 응답 캐시 표출 + 🔴 ApiStatusBar"],
        ["멱등성 키", "X-Request-Id 동일 시 동일 응답 반환 (5분 윈도우)"],
    ],
    col_widths=[5.0, 12.0]
)

add_heading(doc, "8.4 감사 / 로깅", level=2)
for t in [
    "모든 API 호출은 양 측에서 로깅 (request_id, timestamp, source_ip, user, latency, status)",
    "감사 로그 보존 : 최소 1년 (운영기관 정책에 따름)",
    "에러 로그는 별도 SIEM(또는 운영기관 표준)으로 전송",
    "센서 raw 데이터 외부 전송 시 로그에는 sensor_id/통계값만 (raw 값은 별도 스토리지)",
]:
    add_bullet(doc, t)

doc.add_page_break()

# ─────────────── 9. 시퀀스 다이어그램 ───────────────
add_heading(doc, "9. 시퀀스 다이어그램 (이벤트 발생 → 종료 전 흐름)", level=1)

add_heading(doc, "9.1 SYMPTOM 단계 — 자동 enrichment", level=2)
add_code(doc,
    "(SYMPTOM 진입)\n"
    "  플랫폼 ──GET /api/provider/safetia/{scenario_id}──→ 세이프티아\n"
    "  플랫폼 ←──[ 이력 배열 ]────────────────────────── 세이프티아\n"
    "\n"
    "  플랫폼 ──GET /api/provider/kogas/{scenario_id}──→ KOGAS\n"
    "  플랫폼 ←──[ 진단 결과 ]─────────────────────── KOGAS\n"
    "\n"
    "  플랫폼  → SSE: SYMPTOM_ENRICHMENT (kogas_result, safetia_history)"
)

add_heading(doc, "9.2 FAULT 단계 — 위험예측", level=2)
add_code(doc,
    "(FAULT 진입 — EventLog OPEN)\n"
    "  운영자  → 위험예측 모드(M-RSK)에서 [실행] 클릭\n"
    "  플랫폼 ──POST /api/provider/kgs/analyze {sensor_data, params}──→ KGS\n"
    "  플랫폼 ←──[ 영향 분석 배열 ]──────────────────────────────── KGS\n"
    "  플랫폼  → 2D 네트워크 + 3D 컬러링 + HAZOP 텍스트 표출\n"
    "  플랫폼  → SOP 추천 활성화"
)

add_heading(doc, "9.3 SECONDARY_IMPACT 단계 — 시뮬레이션", level=2)
add_code(doc,
    "(SECONDARY_IMPACT 진입)\n"
    "  운영자  → 시뮬레이션 모드(M-SIM)에서 [실행] 클릭\n"
    "  플랫폼 ──POST /api/provider/keti/simulate {scenario_id, params}──→ KETI\n"
    "  플랫폼 ←──[ Option A/B + summary ]───────────────────────── KETI\n"
    "  플랫폼  → 대응안 비교 카드 표출 + 3D 가스확산\n"
    "  운영자  → [Option A 적용]  → SOP 추천 갱신"
)

add_heading(doc, "9.4 RESPONSE 단계 — 보고서 자동생성 + 이력 등록", level=2)
add_code(doc,
    "(RESPONSE 진입 — EventLog CLOSED)\n"
    "  플랫폼  → 자체 로직: report_generator.generateReport(event_id)\n"
    "          (KOGAS·KGS·KETI·세이프티아 데이터 통합 → ReportDocument 생성)\n"
    "\n"
    "  [본개발] 플랫폼 ──POST /v1/history (신규 이력 등록) ──→ 세이프티아\n"
    "          ←── 201 Created { history_id }\n"
    "\n"
    "  [본개발 옵션] 플랫폼 ──POST /v1/feedback (정/오진 라벨) ──→ KOGAS\n"
    "          (KOGAS 모델 재학습 데이터로 활용)"
)

doc.add_page_break()

# ─────────────── 10. 인터페이스 합의 체크리스트 ───────────────
add_heading(doc, "10. 인터페이스 합의 체크리스트", level=1)

add_heading(doc, "10.1 KOGAS 측 확인 사항", level=2)
add_table(doc,
    ["#", "항목", "확정", "비고"],
    [
        ["K1", "fault_code 표준 체계 (FLT-* 외 신규)", "□", ""],
        ["K2", "sensor_evidence 형식 (string vs object)", "□", ""],
        ["K3", "응답 시간 SLA (동기 5초)", "□", ""],
        ["K4", "Push (webhook) 적용 여부", "□", ""],
        ["K5", "센서 데이터 push 주기/범위", "□", ""],
        ["K6", "재학습 데이터 회신 정책", "□", ""],
        ["K7", "동일 설비 다중 동시 진단 처리", "□", ""],
        ["K8", "신뢰도 산출 기준", "□", ""],
        ["K9", "인증 방식 (API Key / OAuth / mTLS)", "□", ""],
        ["K10", "감사 로그 보존 정책", "□", ""],
    ],
    col_widths=[1.0, 9.0, 1.5, 5.5]
)

add_heading(doc, "10.2 KGS 측 확인 사항", level=2)
add_table(doc,
    ["#", "항목", "확정", "비고"],
    [
        ["G1", "impact_type enum 추가 (현 8종 → ?)", "□", ""],
        ["G2", "impact_score 산출 기준", "□", ""],
        ["G3", "predicted_after_sec 정밀도", "□", ""],
        ["G4", "응답 시간 SLA", "□", ""],
        ["G5", "영향 그래프 깊이 (depth)", "□", ""],
        ["G6", "recommended_action 텍스트 표준 / SOP 매핑", "□", ""],
        ["G7", "color_2d / color_3d 위치 (KGS / 플랫폼)", "□", ""],
        ["G8", "센서 데이터 입력 범위", "□", ""],
        ["G9", "인증 방식", "□", ""],
        ["G10", "감사 로그 보존 정책", "□", ""],
    ],
    col_widths=[1.0, 9.0, 1.5, 5.5]
)

add_heading(doc, "10.3 KETI 측 확인 사항", level=2)
add_table(doc,
    ["#", "항목", "확정", "비고"],
    [
        ["E1", "대응안 개수 (A/B 고정 vs N개)", "□", ""],
        ["E2", "동기/비동기 응답 시간 SLA", "□", ""],
        ["E3", "비동기 콜백 URL 등록 방식", "□", ""],
        ["E4", "timeseries_forecast 제공 여부", "□", ""],
        ["E5", "simulation_summary 길이 제한", "□", ""],
        ["E6", "동적 저작도구 연동 워크플로우", "□", ""],
        ["E7", "fault_type enum 표준", "□", ""],
        ["E8", "결과 영구 저장 정책", "□", ""],
        ["E9", "인증 방식", "□", ""],
        ["E10", "감사 로그 보존 정책", "□", ""],
    ],
    col_widths=[1.0, 9.0, 1.5, 5.5]
)

add_heading(doc, "10.4 세이프티아 측 확인 사항", level=2)
add_table(doc,
    ["#", "항목", "확정", "비고"],
    [
        ["S1", "history_type 표준 체계", "□", ""],
        ["S2", "페이지네이션·검색 옵션", "□", ""],
        ["S3", "이력 영구 보관 vs 캐시 정책", "□", ""],
        ["S4", "이벤트 종료 시 자동 기록 적용", "□", ""],
        ["S5", "사고대응 라이브러리 연동", "□", ""],
        ["S6", "첨부파일 endpoint", "□", ""],
        ["S7", "응답 시간 SLA", "□", ""],
        ["S8", "운영자 메모 한글 길이 제한", "□", ""],
        ["S9", "인증 방식", "□", ""],
        ["S10", "감사 로그 보존 정책", "□", ""],
    ],
    col_widths=[1.0, 9.0, 1.5, 5.5]
)

doc.add_page_break()

# ─────────────── 11. 부록 ───────────────
add_heading(doc, "11. 부록", level=1)

add_heading(doc, "11.1 PoC 시나리오 ↔ 외부기관 데이터 매트릭스", level=2)
add_table(doc,
    ["시나리오", "Trigger", "KOGAS", "KGS", "KETI", "세이프티아", "SOP"],
    [
        ["SC-01", "BOG-201", "FLT-01 (75%)", "4건 (88/76/64/58)", "REL/벤트 18min/25min", "BOG-201 이력", "SOP-BOG-TRIP-01"],
        ["SC-02", "PMP-301", "FLT-02 (78%)", "3건", "정지/감량 15/35min", "PMP-301", "SOP-PUMP-CAV-01"],
        ["SC-03", "TK-101", "FLT-03 (81%)", "3건", "압축기 최대/이송 32/22min", "TK-101", "SOP-TANK-PRES-01"],
        ["SC-04", "VAL-601", "FLT-04 (84%)", "3건", "Bypass/펌프정지 12/40min", "VAL-601", "SOP-VAL-STUCK-01"],
        ["SC-05", "PIP-501", "FLT-05 (72%)", "2건", "—", "PIP-501", "SOP-BOG-SURGE-01"],
        ["SC-06", "VAP-401", "FLT-06 (75%)", "2건", "—", "VAP-401", "SOP-VAP-COLD-01"],
        ["SC-07", "TK-101", "FLT-07 (78%)", "1건", "—", "TK-101", "SOP-NORMAL-OPS-01"],
        ["SC-08", "ARM-101", "FLT-ARM-ESD (82%)", "다건", "—", "ARM-101", "SOP-ARM-ESD-01"],
    ],
    col_widths=[1.5, 1.8, 2.4, 2.5, 3.5, 2.0, 3.3]
)

add_heading(doc, "11.2 식별자 / Enum 빠른 참조", level=2)
add_table(doc,
    ["분류", "값"],
    [
        ["EquipmentType", "LH2_CARRIER, LOADING_ARM, STORAGE_TANK, BOG_COMPRESSOR, TRANSFER_PUMP, VAPORIZER, MAIN_PIPE, VALVE_STATION, RELIQUEFIER, SEAWATER_PUMP"],
        ["SensorType", "PRESSURE, TEMPERATURE, FLOW, VIBRATION, CURRENT, LEVEL"],
        ["Severity / risk_level", "INFO, WARNING, CRITICAL, EMERGENCY"],
        ["impact_type (KGS)", "PRIMARY_EVENT, PRESSURE_PROPAGATION, PSV_ACTIVATION_RISK, BACKUP_ACTIVATION, FLOW_INTERRUPTION, PRESSURE_BUILDUP, SUPPLY_SHORTAGE, FLOW_REDUCTION"],
        ["Phase", "NORMAL, SYMPTOM, FAULT, SECONDARY_IMPACT, RESPONSE"],
        ["history_type (세이프티아)", "MAINTENANCE, INSPECTION, INCIDENT, REPLACE, NORMAL_PROFILE"],
        ["fault_type (KETI 요청)", "OVERPRESSURE, CAVITATION, STUCK_VALVE, COLD_BRITTLE, BOG_SURGE, ESD_TRIP, NORMAL_OPS"],
        ["option_*_risk (KETI)", "NONE, LOW, MEDIUM, HIGH"],
        ["return_mode (KETI)", "sync, async"],
        ["SOP Category", "EMERGENCY, SAFETY, ROUTINE, INSPECTION, EVENT_RESPONSE, MAINTENANCE"],
    ],
    col_widths=[3.5, 13.5]
)

add_heading(doc, "11.3 Health Check 통합 응답 예시", level=2)
add_code(doc,
    "GET /api/provider/kogas/health    → 200 { \"status\": \"ok\", \"provider\": \"KOGAS\" }\n"
    "GET /api/provider/kgs/health      → 200 { \"status\": \"ok\", \"provider\": \"KGS\" }\n"
    "GET /api/provider/keti/health     → 200 { \"status\": \"ok\", \"provider\": \"KETI\" }\n"
    "GET /api/provider/safetia/health  → 200 { \"status\": \"ok\", \"provider\": \"SAFETIA\" }\n"
    "\n"
    "비정상 시:\n"
    "  503 { \"status\": \"degraded\" | \"down\", \"provider\": \"<name>\", \"reason\": \"...\" }"
)

add_heading(doc, "11.4 변경 이력", level=2)
add_table(doc,
    ["버전", "일자", "주요 변경", "작성자"],
    [
        ["v1.0", "2026-05-07", "PoC 기반 본개발 인터페이스 합의안 최초 작성", "유엔이"],
    ],
    col_widths=[2.0, 3.0, 9.0, 3.0]
)

add_heading(doc, "11.5 검토자 서명란", level=2)
add_table(doc,
    ["기관", "검토자(직책/성명)", "검토일", "의견 / 서명"],
    [
        ["KOGAS", "", "", ""],
        ["KGS", "", "", ""],
        ["KETI", "", "", ""],
        ["세이프티아", "", "", ""],
        ["운영기관", "", "", ""],
        ["유엔이(작성)", "", "", ""],
    ],
    col_widths=[3.0, 4.5, 3.0, 6.5]
)

# 저장
out_path = "/home/user/UNE/docs/IFD_LH2_DigitalTwin_v1.0.docx"
doc.save(out_path)
print(f"[OK] saved: {out_path}")
