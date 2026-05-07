# -*- coding: utf-8 -*-
"""
LH2 액화수소 인수기지 디지털 트윈 자율안전관리 플랫폼
상세 요구사항 정의서 (Requirements Specification) docx 생성기

작성 근거:
- /home/user/UNE/CLAUDE.md (POC 단일 개발 가이드)
- /home/user/UNE/apps/api/src/** (Express + Prisma 백엔드 구현체)
- /home/user/UNE/apps/web/src/** (Next.js 14 프론트엔드 구현체)
- /home/user/UNE/seed/*.json (28개 시드 데이터)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_default_font(doc, name="맑은 고딕", size=10):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.runs[0] if p.runs else p.add_run("")
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.name = "맑은 고딕"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "맑은 고딕"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], "1F3A5F")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "맑은 고딕"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


# ────────────────────────────────────────────────────────────────────
# 문서 생성 시작
# ────────────────────────────────────────────────────────────────────
doc = Document()
set_default_font(doc)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ─────────────── 표지 ───────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("액화수소 인수기지\n디지털 트윈 자율안전관리 플랫폼")
r.bold = True
r.font.size = Pt(24)
r.font.name = "맑은 고딕"
rPr = r._element.get_or_add_rPr()
rF = OxmlElement("w:rFonts"); rPr.append(rF)
rF.set(qn("w:eastAsia"), "맑은 고딕")
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(20)
r = sub.add_run("상세 요구사항 정의서 (Requirements Specification)")
r.bold = True
r.font.size = Pt(18)
r.font.name = "맑은 고딕"
rPr = r._element.get_or_add_rPr(); rF = OxmlElement("w:rFonts"); rPr.append(rF)
rF.set(qn("w:eastAsia"), "맑은 고딕")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(180)
for line in [
    "문서 버전 : v1.0 (POC 역설계 → 본개발 기준안)",
    "작성 일자 : 2026-05-07",
    "작성 주체 : 유엔이 (플랫폼 개발사)",
    "검토 대상 : KOGAS · KGS · KETI · 세이프티아 · 운영기관",
    "기준 문서 : CLAUDE.md, FUNC_SPEC_POC_v5.md, POC 구현 코드베이스",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(line)
    rr.font.size = Pt(11)
    rr.font.name = "맑은 고딕"
    rPr = rr._element.get_or_add_rPr(); rF = OxmlElement("w:rFonts"); rPr.append(rF)
    rF.set(qn("w:eastAsia"), "맑은 고딕")

doc.add_page_break()

# ─────────────── 문서 개요 ───────────────
add_heading(doc, "0. 문서 개요", level=1)

add_heading(doc, "0.1 문서의 목적", level=2)
add_para(doc,
    "본 문서는 「액화수소 인수기지 디지털 트윈 자율안전관리 플랫폼」 본개발(이하 \"본개발\")의 "
    "기능·비기능 요구사항을 정의한다. 본개발에 앞서 수행된 PoC(Proof of Concept) 결과물의 "
    "구현 범위와 동작 방식을 역설계(reverse-engineering)하여 정리한 사양이며, 참여기관(KOGAS, "
    "KGS, KETI, 세이프티아)이 본 문서를 검토하고 모니터링 데이터·상세 기능을 확정한 후 본개발에 "
    "착수하기 위한 기준안의 역할을 수행한다."
)
add_para(doc,
    "본 문서와 함께 배포되는 「기관별 데이터 인터페이스 정의서(IFD)」는 외부기관 서버와 플랫폼 "
    "간의 데이터 교환 인터페이스만을 별도로 다루며, 본 요구사항 정의서는 플랫폼 자체 기능에 "
    "초점을 맞춘다."
)

add_heading(doc, "0.2 적용 범위 및 제외 범위", level=2)
add_para(doc, "(1) 본 요구사항 정의서가 다루는 범위", bold=True)
for t in [
    "플랫폼이 직접 개발·운영하는 기능 (3D/2D 통합 모니터링, 알람·이벤트 관리, 센서수집·표출, 외부기관 결과 시각화, SOP 저작/실행, 보고서 자동생성, 설정관리, 시나리오 에뮬레이터)",
    "플랫폼과 외부기관 시스템 간의 연동 인터페이스 정의(요약). 상세 인터페이스는 별첨 IFD 참조",
    "본개발 단계의 비기능 요구사항(성능, 보안, 배포, 운영)",
    "PoC에서 확인된 데이터 모델(DB 스키마)과 외부기관 연동 데이터의 사전 합의안",
]:
    add_bullet(doc, t)

add_para(doc, "(2) 본 문서가 다루지 않는 범위 (별도 기관 책임)", bold=True)
for t in [
    "KOGAS : 이상탐지/고장인지 AI 모델의 학습 알고리즘 및 내부 구조",
    "KGS : 상호영향 위험예측 엔진의 분석 알고리즘 및 위험도 산출 로직",
    "KETI : 시뮬레이션 엔진/하이브리드 모델링/동적 저작도구의 내부 구조",
    "세이프티아 : 사고대응/이력관리 시스템의 데이터베이스 설계",
    "기관 내부 운영 정책 및 보안 통제(별도 기관 책임)",
]:
    add_bullet(doc, t)

add_heading(doc, "0.3 용어 정의", level=2)
add_table(doc,
    ["용어", "정의"],
    [
        ["플랫폼", "유엔이가 개발·운영하는 디지털 트윈 자율안전관리 플랫폼 (Frontend+Backend+DB)"],
        ["외부기관", "KOGAS, KGS, KETI, 세이프티아 (각자 자체 서버 운영)"],
        ["Mock Provider", "PoC 단계에서 외부기관 서버 부재 시 동일 응답 스펙으로 대체하는 모듈 (본개발 시 실제 API로 교체)"],
        ["EventContext", "이벤트 발생 시 플랫폼 전역에 공유되는 상황 정보 (시나리오ID, 트리거 설비, 영향 설비, 단계, 외부기관 결과)"],
        ["Phase", "시나리오 진행 단계 (NORMAL → SYMPTOM → FAULT → SECONDARY_IMPACT → RESPONSE)"],
        ["SOP", "Standard Operating Procedure (표준 작업 절차서, 단계형 체크리스트)"],
        ["HAZOP", "Hazard and Operability Study (위험성평가)"],
        ["GLB", "glTF Binary (3D 모델 파일 포맷, Draco 압축)"],
        ["SSE", "Server-Sent Events (서버→클라이언트 단방향 실시간 스트림)"],
    ],
    col_widths=[3.5, 13.0]
)

add_heading(doc, "0.4 참조 문서", level=2)
for t in [
    "CLAUDE.md — POC 단일 개발 가이드 (본 문서의 직접 근거)",
    "FUNC_SPEC_POC_v5.md — POC 기능명세서",
    "DB_SCHEMA_POC_v2.md — DB 스키마 (Prisma schema로 단일화됨)",
    "/seed/seed_*.json — 28개 시드 데이터 파일",
    "/apps/api/src/** — Express + Prisma 백엔드 구현체",
    "/apps/web/src/** — Next.js 14 + Three.js 프론트엔드 구현체",
    "「기관별 데이터 인터페이스 정의서(IFD) v1.0」 — 별첨",
]:
    add_bullet(doc, t)

doc.add_page_break()

# ─────────────── 1. 시스템 개요 ───────────────
add_heading(doc, "1. 시스템 개요", level=1)

add_heading(doc, "1.1 비전과 목표", level=2)
add_para(doc,
    "본 플랫폼은 액화수소 인수기지의 핵심설비(저장탱크, BOG 압축기, 이송펌프, 기화기, 메인배관, "
    "벤트스택, 재액화기, 운반선, 로딩암 등)를 3D/2D 디지털 트윈으로 구축하여, 평시에는 상시 "
    "모니터링과 이상감지를, 비상시에는 위험예측·시뮬레이션·대응절차 실행·보고서 자동생성까지를 "
    "단일 화면에서 수행할 수 있도록 한다."
)
add_para(doc,
    "PoC 단계에서는 8종 시나리오(SC-01 ~ SC-08) 기반 시뮬레이션으로 전 기능을 시현하였으며, "
    "본개발 단계에서는 실데이터 연동·실 외부기관 API 연동·다중 사용자 동시접속·운영기관 보안정책 "
    "준수까지 확장한다."
)

add_heading(doc, "1.2 시스템 구성 (논리)", level=2)
add_code_block(doc,
    "[운영자 / 외부기관 검토자 브라우저]\n"
    "          ↕ HTTPS\n"
    "[Frontend — Next.js 14 (Vercel)]\n"
    "          ↕ REST API / SSE\n"
    "[Backend — Node.js + Express + Prisma (Railway)]\n"
    "    ↕                      ↕                      ↕\n"
    "[PostgreSQL]   [Cloudflare R2 (GLB)]   [외부기관 서버]\n"
    "                                          - KOGAS  : 이상탐지/고장인지 AI\n"
    "                                          - KGS    : 상호영향 위험예측\n"
    "                                          - KETI   : 시뮬레이션 엔진\n"
    "                                          - 세이프티아 : 사고대응/이력 DB"
)

add_heading(doc, "1.3 외부기관 책임 분담 원칙 (중요)", level=2)
add_para(doc,
    "본 플랫폼은 외부기관 시스템과 \"느슨한 결합(loose coupling)\" 구조로 연동된다. 외부기관은 "
    "각자 서버를 운영하며, 플랫폼은 표준화된 REST API로만 통신한다. PoC에서는 외부기관 서버 "
    "부재로 Mock Provider가 동일한 응답 스펙으로 대체하였으나, 본개발 단계에서는 동일 스펙을 "
    "유지한 채 실제 기관 서버로 endpoint를 교체한다."
)
add_table(doc,
    ["기관", "담당 영역", "PoC 단계", "본개발 단계"],
    [
        ["KOGAS", "이상탐지/고장인지 AI 모델 운영", "Mock Provider가 scenario_id 기준 진단결과 JSON 반환", "KOGAS 자체 서버 → REST API 연동"],
        ["KGS", "상호영향 위험예측 엔진 운영", "Mock Provider가 scenario_id 기준 영향설비/위험도 JSON 반환", "KGS 자체 서버 → REST API 연동"],
        ["KETI", "시뮬레이션 엔진 + 동적 저작도구 운영", "Mock Provider가 scenario_id 기준 시뮬레이션 결과 JSON 반환", "KETI 자체 서버 → REST API 연동"],
        ["세이프티아", "사고대응/이력관리 시스템 운영", "Mock Provider가 scenario_id 기준 이력 JSON 반환", "세이프티아 자체 서버 → REST API 연동"],
        ["유엔이 (플랫폼)", "통합 모니터링·시각화·SOP·보고서·설정·에뮬레이터", "전 기능 직접 구현 (Mock Provider 4종 포함)", "전 기능 + 운영기관 보안정책 적용"],
    ],
    col_widths=[2.8, 5.0, 4.3, 4.3]
)

add_heading(doc, "1.4 사용자 유형 (Persona)", level=2)
add_table(doc,
    ["역할", "주요 권한 / 사용 모드", "PoC", "본개발"],
    [
        ["운영자(중앙제어실)", "M-MON 상시 모니터링, 이벤트 인지, SOP 실행", "단일 역할", "RBAC 적용"],
        ["현장 점검원", "M-ANO 설비 상태감시, 현장 조치 후 SOP 체크", "단일 역할", "모바일 최적화 + RBAC"],
        ["안전관리자", "M-RSK, M-SIM 위험예측·시뮬레이션, 대응안 결정", "단일 역할", "RBAC + 결재라인"],
        ["분석/설계자", "M-HIS 이력분석, 보고서 검토, P-SET 설정", "단일 역할", "RBAC + 감사로그"],
        ["기관 검토자", "POC 화면 체험 + 본개발 인터페이스 검토", "PoC 한정", "RBAC 외부 게스트"],
    ],
    col_widths=[3.5, 7.0, 3.0, 3.0]
)

doc.add_page_break()

# ─────────────── 2. 기능 요구사항 — 모드 구조 ───────────────
add_heading(doc, "2. 기능 요구사항 — 모드 체계", level=1)

add_heading(doc, "2.1 모드 구성 (총 8개)", level=2)
add_para(doc, "플랫폼은 6개의 운영 모드와 2개의 보조 페이지로 구성된다.", bold=True)
add_table(doc,
    ["코드", "모드/페이지명", "정의", "주요 외부기관 연동"],
    [
        ["M-MON", "기본 모니터링", "3D/2D 통합 공간 보기, 설비 상태, 센서값, 알람 확인 (상시)", "—"],
        ["M-ANO", "설비 상태감시 / 이상탐지", "이상진단·센서추세·이상탐지 점수 표출", "KOGAS"],
        ["M-RSK", "위험예측", "2D 영향 네트워크 + 3D 컬러링 + HAZOP 텍스트 통합", "KGS, HAZOP"],
        ["M-SIM", "시뮬레이션 / 의사결정지원", "시뮬레이션 결과·대응안 비교, 가스확산 시각화", "KETI, KGS"],
        ["M-HIS", "이력조회 / 분석", "정비/사고/점검 이력, 추세분석", "세이프티아"],
        ["M-SOP", "디지털 SOP", "SOP 실행(메인) + 저작/편집 + 실행이력 + 휴지통", "—"],
        ["P-SET", "설정", "센서 메타데이터, 임계치, 운영정책 관리", "—"],
        ["P-RPT", "보고서", "이벤트 조치 보고서 자동생성·편집·제출", "전 기관 데이터 통합"],
    ],
    col_widths=[1.8, 3.5, 8.0, 3.5]
)

add_heading(doc, "2.2 모드 간 연계 원칙 (EventContext)", level=2)
add_para(doc,
    "이벤트가 발생하면 그 정보(EventContext)는 모든 모드에 걸쳐 공유되며, 모드 전환 시에도 "
    "동일한 상황 정보가 유지된다. 운영자가 \"이상탐지 → 위험예측 → 시뮬레이션 → SOP 실행 → "
    "보고서 자동생성\"의 흐름을 끊김 없이 수행할 수 있어야 한다."
)
add_code_block(doc,
    "EventContext = {\n"
    "  event_id, scenario_id, trigger_equipment_id,\n"
    "  affected_equipment_ids[], severity, current_phase,\n"
    "  hazop_id,\n"
    "  kogas_result?, kgs_results?[], keti_result?, safetia_history?[],\n"
    "  recommended_sops?[]\n"
    "}"
)

add_heading(doc, "2.3 SOP 듀얼 UI 원칙", level=2)
add_para(doc,
    "운영 모드에서 이벤트 팝업의 [SOP] 버튼을 클릭하면 두 가지 동작을 사용자가 선택할 수 있다. "
    "두 UI는 동일 컴포넌트의 크기 변형이며, 진행률·체크상태·메모는 양쪽에서 일관되게 동기화된다."
)
add_table(doc,
    ["방식", "동작", "용도"],
    [
        ["[SOP 팝업 실행]", "현재 모드 유지 + 우측에 축소 SOP 패널(폭 400px)", "조감을 유지하며 절차 수행"],
        ["[SOP 모드로 이동]", "M-SOP 전체 화면으로 전환", "체크리스트 집중 수행"],
    ],
    col_widths=[3.5, 8.0, 5.0]
)

doc.add_page_break()

# ─────────────── 3. 모드별 상세 요구사항 ───────────────
add_heading(doc, "3. 모드별 상세 요구사항", level=1)

# 3.1 M-MON
add_heading(doc, "3.1 M-MON — 기본 모니터링", level=2)
add_para(doc, "(1) 화면 구성 (데스크톱 기준)", bold=True)
add_table(doc,
    ["영역", "비율", "구성요소"],
    [
        ["좌측 공정 흐름 패널", "15%", "4단계 공정(하역→저장·BOG→이송→기화·송출), 단계별 설비 그룹화, 이상 시 단계 박스 강조 + 화살표 라벨에 이상 센서값 표출"],
        ["중앙 3D 뷰어", "65%", "h2.glb 테스트베드 모델, 설비 컬러링, POI 마커, 카메라 프리셋(12종), 카메라 시점 저장/복원"],
        ["우측 정보 패널", "20%", "선택 설비 정보, 센서 현재값 4종, 알람 이력, 모드 전환 버튼(이상탐지/위험예측/시뮬레이션/SOP/이력조회)"],
        ["하단 KPI 대시보드", "전체", "9개 핵심설비의 대표 센서값과 상태(🔵정상/🟡경고/🔴위험), 카드 클릭 시 카메라 이동"],
    ],
    col_widths=[4.0, 2.0, 11.0]
)
add_para(doc, "(2) 핵심 기능", bold=True)
for t in [
    "공정 흐름 패널은 4단계 (하역 → 저장·BOG → 이송 → 기화·송출) 의 카드형 레이아웃으로 구성하며, 단계 간 LH2 액체/BOG 가스 흐름을 화살표로 표현한다.",
    "이상 발생 시: trigger 설비가 속한 단계 박스에 색상(주황/적색) + 펄스 애니메이션, 화살표 라벨에 이상 센서값을 조건부 표시.",
    "BOG 순환 루프(저장탱크 → BOG 발생 → BOG 압축기 → 재액화기 → 저장탱크)와 벤트스택 비상 경로를 시각적으로 표현.",
    "설비 클릭 시: 카메라 이동(0.8초 GSAP 트랜지션) + 우측 정보 패널 갱신 + 공정 흐름 패널의 동일 설비 하이라이트.",
    "이벤트 팝업: severity, phase, 영향설비, KOGAS 진단 요약, 권고 SOP, 모드 전환 5개 버튼 + SOP 듀얼 실행 버튼 포함.",
    "에뮬레이터 미실행(상시 모니터링) 시: AmbientProvider가 95초 가상 운전 사이클(SHIP_APPROACH→ARM_CONNECT→UNLOADING→STORAGE_PROCESS→TRANSFER→IDLE)로 더미 센서값 송출.",
]:
    add_bullet(doc, t)

# 3.2 M-ANO
add_heading(doc, "3.2 M-ANO — 설비 상태감시 / 이상탐지", level=2)
add_para(doc, "(1) 화면 구성", bold=True)
add_table(doc,
    ["영역", "구성요소"],
    [
        ["설비 선택 탭", "9개 핵심설비를 수평 탭으로 전환 (TK-101, TK-102, BOG-201, PMP-301, VAP-401, ARM-101 등)"],
        ["좌·우 센서차트 (각 3-4개)", "선택 설비의 주요 센서 시계열 + 임계선(경고/위험) + 이상구간 적색 배경"],
        ["중앙 3D 설비 뷰어", "PMP-301: secondary_pump.glb X-ray 모드 (impeller_stage_03/04 적색, shaft 주황). 그 외 설비: 격리 뷰(IsolatedEquipmentModel) + 카메라 줌"],
        ["하단 이상탐지 그래프", "(좌) 24h 타임라인 실측 vs 학습값 비교 차트, (중) 시간별 상세 테이블(기준값/학습값/오차), (우) KOGAS 비교/이상탐지 구간 + 진단결과 텍스트"],
        ["KOGAS 진단 결과 바", "API 연결상태, 고장명, 확신도, 의심부위, 고장코드 + 모드전환 4개 버튼"],
    ],
    col_widths=[5.0, 12.0]
)
add_para(doc, "(2) 외부 연동", bold=True)
for t in [
    "GET /api/provider/kogas/{scenario_id} → KOGAS 진단결과(fault_code, fault_name, diagnosis_confidence, suspected_part, sensor_evidence) 표출",
    "GET /api/provider/kogas/health → 연결정상 표시 (상단 ApiStatusBar)",
    "본개발 단계에서는 KOGAS의 실시간 push 형태로 변경 가능 (인터페이스 정의서 §4 참조)",
]:
    add_bullet(doc, t)

# 3.3 M-RSK
add_heading(doc, "3.3 M-RSK — 위험예측", level=2)
add_para(doc, "(1) 화면 구성 (3분할 + 입력 패널)", bold=True)
add_table(doc,
    ["영역", "구성요소"],
    [
        ["분석 입력 패널 (상단)", "이벤트 연계 모드(EventContext 자동 채움) / 수동 모드(설비·시간·KGS 파라미터 입력) + KGS 연결 상태 + [위험예측 실행] 버튼"],
        ["좌측 — 2D 영향 네트워크 (30%)", "react-flow 또는 Canvas 기반 노드-엣지 그래프. 노드 크기 = impact_score, 엣지 = 영향 전파, zone 영향 막대 표시"],
        ["중앙 — 3D/2D 뷰어 (40%)", "trigger 적색·affected 황색 컬러링, 영향 전파 점선 애니메이션, 시간축 슬라이더(0~60분)로 3D 컬러링 단계적 변화, 가스확산 시각화(M-SIM 공유)"],
        ["우측 — HAZOP + 상세 (30%)", "위험도예측/피해범위예측 탭, HAZOP 원인·이벤트·위험·예방·비상조치, 권고조치, 연계 SOP 듀얼 실행 버튼"],
    ],
    col_widths=[5.0, 12.0]
)
add_para(doc, "(2) 외부 연동", bold=True)
for t in [
    "POST /api/provider/kgs/analyze {scenario_id, equipment_id, sensor_data, params} → KGS 위험예측 결과 배열 반환",
    "GET /api/provider/kgs/{scenario_id} → KGS 결과 일괄 조회",
    "GET /api/hazop/{scenario_id} → HAZOP 상세 (node, deviation, cause, hazard_scenario, preventive_action, emergency_response)",
    "이벤트 연계 시: 센서값 + 설비ID 자동 전송 (사용자 입력 없음)",
    "수동 실행 시: 설비/분석시간/압력초과율/온도편차 등 KGS 파라미터를 사용자가 입력 후 [실행]",
]:
    add_bullet(doc, t)

# 3.4 M-SIM
add_heading(doc, "3.4 M-SIM — 시뮬레이션 / 의사결정지원", level=2)
add_para(doc, "(1) 화면 구성", bold=True)
add_table(doc,
    ["영역", "구성요소"],
    [
        ["탭 (이벤트 연계 / 수동 실행)", "이벤트 연계: EventContext 자동 채움 / 수동: 시나리오·이상유형·초기압력·지속시간·온도편차·유량변화 슬라이더"],
        ["3D 시뮬레이션 뷰어 (50%)", "GasDispersion(가스확산 파티클), trigger 적색 점멸, affected 점진 컬러링, 영향반경 반투명 원형 오버레이, 타임라인 스크러버"],
        ["결과 패널 (50%)", "(상) KGS 위험영향: 설비별 점수+위험도+예상시간 / (중) KETI 대응안 A/B 카드: 설명+안정화시간 바+세부정보+[적용▶] 버튼 / (하) 시뮬레이션 요약 + SOP/보고서 연계 버튼"],
    ],
    col_widths=[5.0, 12.0]
)
add_para(doc, "(2) 외부 연동", bold=True)
for t in [
    "POST /api/provider/keti/simulate {scenario_id} → KETI 시뮬레이션 결과 (recommended_option_a/b, option_a/b_stabilization_min, option_a/b_risk, option_a/b_detail, expected_stabilization_min, simulation_summary)",
    "POST /api/provider/kgs/analyze 병행 호출 → 영향설비 컬러링 동기화",
    "대응안 [적용▶] 클릭 시: 3D에서 영향범위 축소 애니메이션 + SOP 추천 갱신",
]:
    add_bullet(doc, t)

# 3.5 M-HIS
add_heading(doc, "3.5 M-HIS — 이력조회 / 분석", level=2)
add_para(doc, "(1) 화면 구성", bold=True)
add_table(doc,
    ["영역", "구성요소"],
    [
        ["좌측 필터 패널", "설비 체크박스(9개), 기간(최근1개월/3개월/1년/전체), 유형(정비/점검/교체/사고)"],
        ["메인 테이블", "5열: 구분 / 설비 / 최근점검일 / 사고이력 / 연계 SOP"],
        ["상세 패널 (하단)", "선택된 이력의 일자·시나리오·유형·요약·운영자 메모·연계 SOP + [위험예측][SOP] 모드 전환 버튼"],
    ],
    col_widths=[5.0, 12.0]
)
add_para(doc, "(2) 외부 연동", bold=True)
for t in [
    "GET /api/provider/safetia/{scenario_id} → 세이프티아 이력(history_id, last_maintenance_date, past_incident_summary, linked_sop_id, operator_note)",
    "본개발에서는 페이지네이션·필터·기간 검색을 외부기관 서버에서 직접 처리하도록 인터페이스 확장",
]:
    add_bullet(doc, t)

# 3.6 M-SOP
add_heading(doc, "3.6 M-SOP — 디지털 SOP", level=2)
add_para(doc, "(1) 탭 구성 (4개)", bold=True)
add_table(doc,
    ["탭", "주요 기능"],
    [
        ["실행 (메인)", "SOP 목록 → SOP 플로우차트 실행 UI(텍스트/체크박스/결정 분기 노드) → 진행률 → 메모 → [실행완료][상황전파]"],
        ["저작 / 편집", "SOP 신규 생성·수정. 단계 추가/삭제/순서 변경, 카테고리 색상, 대상 공간/설비, 우선순위(심각/경계/주의/관심), 카메라 프리셋, 팝업 템플릿"],
        ["실행 이력", "SopExecutionLog 조회 (event_id, scenario_id, sop_id, executor_role, checked_steps, started/ended_at, memo)"],
        ["휴지통", "soft-delete된 SOP를 30일간 보관, 잔여일수 표시, [복원]/[영구삭제]"],
    ],
    col_widths=[3.5, 14.0]
)
add_para(doc, "(2) SOP 추천 로직", bold=True)
add_code_block(doc,
    "이벤트 발생 시:\n"
    "  1) event.trigger_equipment_id → equipment_master.zone_id 조회\n"
    "  2) sop_equipment_map에서:\n"
    "     a) equipment_id 일치 + event_severity ≥ event_severity_min  → 1차 후보\n"
    "     b) zone_id 일치                                              → 2차 후보\n"
    "  3) is_primary=true 우선 + sort_order ASC 정렬\n"
    "  4) 0건 시 SOP-GENERIC-INSPECT-01 (fallback)\n"
    "  5) 대표 SOP 1건 + 관련 SOP 전체 목록 반환\n"
)

# 3.7 P-SET
add_heading(doc, "3.7 P-SET — 설정", level=2)
add_para(doc, "(1) 탭 구성 (3개)", bold=True)
add_table(doc,
    ["탭", "주요 기능"],
    [
        ["센서 메타데이터", "ID/이름/유형/설비/단위/샘플링주기(초)/활성여부 인라인 편집 + [저장]"],
        ["임계치 관리", "설비 드롭다운 → 센서별 normal_value, warning_low/high, critical_low/high 인라인 편집 + [기본값 복원][저장]"],
        ["운영정책", "SOP 자동팝업 ON/OFF, 자동 보고서 초안 ON/OFF, Missing Data Timeout, 기본 샘플링 주기 등"],
    ],
    col_widths=[4.0, 13.0]
)

# 3.8 P-RPT
add_heading(doc, "3.8 P-RPT — 보고서", level=2)
add_para(doc, "(1) 화면 구성", bold=True)
add_table(doc,
    ["영역", "구성요소"],
    [
        ["좌측 보고서 목록", "RPT-* 목록, [+ 생성] 버튼 → 이벤트 선택 모달, 일괄 삭제 체크박스"],
        ["우측 보고서 상세", "(자동수집) 이벤트 개요, 트리거 설비, KOGAS 진단, KGS 영향분석, KETI 권고안, 이력 요약, SOP 수행이력 / (관리자 작성) 의견·후속조치 + [저장][제출][PDF][삭제]"],
    ],
    col_widths=[5.0, 12.0]
)
add_para(doc, "(2) 자동생성 로직", bold=True)
add_code_block(doc,
    "이벤트 RESPONSE 단계 진입(또는 CLOSED) 시:\n"
    "  1) event_log → 이벤트 요약\n"
    "  2) mock_kogas_result(scenario_id) → 진단 요약\n"
    "  3) mock_kgs_result(scenario_id) → 영향설비/위험도\n"
    "  4) mock_keti_result(scenario_id) → 권고안\n"
    "  5) mock_safetia_history(scenario_id) → 이력\n"
    "  6) sop_execution_log(event_id) → 수행이력\n"
    "  7) report_template(RPT-TPL-001) 기반으로 generated_summary 조합\n"
    "  8) report_document INSERT (status='DRAFT')"
)

doc.add_page_break()

# ─────────────── 4. 데이터 요구사항 ───────────────
add_heading(doc, "4. 데이터 요구사항", level=1)

add_heading(doc, "4.1 마스터 데이터 (정적 — 본개발 초기 일괄 적재)", level=2)
add_table(doc,
    ["데이터", "건수(POC)", "본개발 확장 방향"],
    [
        ["Zone (공간/구역)", "8", "실제 인수기지 도면에 따라 확장 (방화구역, ESD 구역 추가 등)"],
        ["Equipment (설비)", "12 (PoC 핵심)", "전수 설비로 확장 + 외부기관 표준 ID 체계 연동"],
        ["Sensor (센서)", "36 (PoC 핵심)", "실제 P&ID 기반 전수 등록 + tag-naming 표준화"],
        ["SensorThreshold (임계치)", "37", "운영기관 기준치 + 계절·운전조건별 동적 임계치"],
        ["HAZOP", "8", "전 시나리오·전 노드로 확장, KGS와 동기화"],
        ["SopCatalog", "9", "운영기관 보유 SOP 전수 등록 + 버전 관리 + 결재 라인"],
    ],
    col_widths=[5.0, 3.0, 9.0]
)

add_heading(doc, "4.2 운영 데이터 (동적 — 실시간/이벤트 기반)", level=2)
add_table(doc,
    ["데이터", "발생 시점", "본개발 확장 방향"],
    [
        ["EventLog", "센서 임계치 초과 / 외부기관 이상감지 시", "다중 이벤트 동시 처리 + 이벤트 통합 (correlation)"],
        ["SopExecutionLog", "운영자가 SOP 실행 시", "단계별 첨부파일·사진·서명 추가, 결재 라인 적용"],
        ["ReportDocument", "이벤트 종료 시 자동, 수동 [생성]", "PDF 출력, 결재라인, 외부기관 회람"],
        ["SensorTimeSeries", "5초 주기 (기본) / 설비별 가변", "PoC: JSON 파일, 본개발: TSDB(InfluxDB/TimescaleDB) 또는 실시간 스트림"],
    ],
    col_widths=[5.0, 5.0, 7.0]
)

add_heading(doc, "4.3 외부기관 결과 데이터 (Mock → Real)", level=2)
add_para(doc, "PoC에서는 \"Mock*\" 테이블이 scenario_id 기반 응답을 보관한다. 본개발에서는 외부기관 서버가 동일 스펙으로 응답하며, 플랫폼은 결과를 캐시하거나 영구 저장한다 (정책 결정 필요).", bold=False)
add_table(doc,
    ["데이터", "PoC 보관", "본개발 정책 옵션"],
    [
        ["KOGAS 진단", "MockKogasResult (Pre-seeded)", "(A) 캐시만 / (B) 영구 저장 + 재학습 데이터로 활용"],
        ["KGS 영향분석", "MockKgsResult (Pre-seeded)", "(A) 캐시만 / (B) 사후 분석용 영구 저장"],
        ["KETI 시뮬레이션", "MockKetiResult (Pre-seeded)", "(A) 호출 시점 1회 / (B) 영구 저장 + 사례 DB"],
        ["세이프티아 이력", "MockSafetiaHistory (Pre-seeded)", "(A) 매 호출 / (B) 캐시(TTL) (C) 부분 동기화"],
    ],
    col_widths=[4.0, 5.0, 8.0]
)

add_heading(doc, "4.4 본 플랫폼 DB 핵심 모델 (Prisma — 16개)", level=2)
add_table(doc,
    ["모델", "용도", "주요 PK/FK"],
    [
        ["Zone", "공간/구역 마스터", "zone_id (PK)"],
        ["EquipmentMaster", "설비 마스터", "equipment_id (PK), zone_id (FK)"],
        ["SensorMaster", "센서 마스터", "sensor_id (PK), equipment_id (FK)"],
        ["SensorThreshold", "센서 임계치", "sensor_id (PK/FK)"],
        ["ScenarioMaster", "시나리오 정의", "scenario_id (PK), trigger_equipment_id (FK)"],
        ["HazopMaster", "HAZOP 분석", "hazop_id (PK), scenario_id (FK), equipment_id (FK)"],
        ["EventLog", "이벤트 로그", "event_id (PK), scenario_id, trigger_equipment_id"],
        ["SopCatalog", "SOP 카탈로그", "sop_id (PK), linked_hazop_id (FK)"],
        ["SopEquipmentMap", "SOP-설비 매핑", "map_id (PK), sop_id (FK), equipment_id (FK)"],
        ["SopExecutionLog", "SOP 실행 로그", "execution_id (PK), event_id (FK), sop_id (FK)"],
        ["ReportDocument", "보고서", "report_id (PK), event_id (FK), template_id"],
        ["SettingsMetadata", "설정", "setting_id (PK), setting_key (unique)"],
        ["MockKogasResult", "KOGAS Mock 결과", "request_id (PK), scenario_id, target_equipment_id"],
        ["MockKgsResult", "KGS Mock 결과", "analysis_id (PK), scenario_id, trigger/affected_equipment_id"],
        ["MockKetiResult", "KETI Mock 결과", "simulation_id (PK), scenario_id, trigger_equipment_id"],
        ["MockSafetiaHistory", "Safetia Mock 이력", "history_id (PK), scenario_id, equipment_id"],
    ],
    col_widths=[5.0, 4.5, 7.5]
)

doc.add_page_break()

# ─────────────── 5. 백엔드 API ───────────────
add_heading(doc, "5. 백엔드 API 요구사항 (REST + SSE)", level=1)
add_para(doc,
    "PoC에서 구현된 67개 엔드포인트의 분류이다. 본개발에서는 API 명세는 동일 유지하고, "
    "인증(JWT/OAuth2), 권한(RBAC), Rate Limiting, 감사 로그를 추가한다."
)

add_heading(doc, "5.1 시나리오 / 에뮬레이터", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/scenarios", "시나리오 목록"],
        ["GET", "/api/scenarios/:id", "시나리오 상세 (phases, hazop_id 포함)"],
        ["POST", "/api/emulator/start", "시나리오 시작 {scenario_id, speed}"],
        ["POST", "/api/emulator/stop", "시나리오 중지"],
        ["POST", "/api/emulator/pause", "일시정지"],
        ["POST", "/api/emulator/resume", "재개"],
        ["GET", "/api/emulator/status", "현재 상태 {running, paused, scenario_id, elapsed_sec, phase, speed, total_duration}"],
        ["GET", "/api/emulator/stream", "SSE 스트림 (CONNECTED, SENSOR_UPDATE, ALARM, PHASE_CHANGE, SYMPTOM_ENRICHMENT, EVENT_CREATE, EVENT_CLOSED, REPORT_GENERATED, SCENARIO_END)"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.2 설비 / 센서 / 공간", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/equipment", "설비 목록"],
        ["GET", "/api/equipment/:id", "설비 상세 (sensors, zone 포함)"],
        ["GET", "/api/equipment/:id/sensors", "설비별 센서 + 임계치"],
        ["GET", "/api/zones", "공간 목록 + 소속 설비"],
        ["GET", "/api/sensors/:id/timeseries", "시계열 (?scenario_id, from, to)"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.3 이벤트 / 알람", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/events", "이벤트 목록 (?status, severity, scenario_id)"],
        ["GET", "/api/events/:id", "이벤트 상세 + KOGAS/KGS/KETI/세이프티아 enrichment"],
        ["PATCH", "/api/events/:id", "상태/심각도/요약 변경"],
        ["GET", "/api/events/stream", "(예약) 이벤트 전용 SSE"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.4 외부기관 Provider (Mock → Real)", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/provider/{kogas|kgs|keti|safetia}/health", "기관별 연결상태 (200/503)"],
        ["GET", "/api/provider/kogas/:scenario_id", "KOGAS 진단 결과"],
        ["GET", "/api/provider/kgs/:scenario_id", "KGS 영향분석 결과(배열)"],
        ["POST", "/api/provider/kgs/analyze", "KGS 위험예측 요청"],
        ["GET", "/api/provider/keti/:scenario_id", "KETI 시뮬레이션 결과"],
        ["POST", "/api/provider/keti/simulate", "KETI 시뮬레이션 요청"],
        ["GET", "/api/provider/safetia/:scenario_id", "세이프티아 이력"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.5 HAZOP", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/hazop", "전체 HAZOP 목록"],
        ["GET", "/api/hazop/:scenario_id", "시나리오별 HAZOP 상세"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.6 SOP", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/sop", "SOP 목록 (?category, equipment_id, status)"],
        ["GET", "/api/sop/trash", "휴지통 (소프트 삭제, 30일 유효)"],
        ["GET", "/api/sop/executions", "실행이력 (?event_id, scenario_id)"],
        ["GET", "/api/sop/recommend", "이벤트 기반 SOP 추천 (?event_id, equipment_id, severity)"],
        ["GET", "/api/sop/:id", "SOP 상세"],
        ["POST", "/api/sop", "SOP 신규 생성"],
        ["PUT", "/api/sop/:id", "SOP 수정"],
        ["DELETE", "/api/sop/:id", "SOP 소프트 삭제"],
        ["POST", "/api/sop/:id/restore", "휴지통에서 복원"],
        ["DELETE", "/api/sop/:id/permanent", "영구 삭제"],
        ["POST", "/api/sop/:id/execute", "SOP 실행 시작"],
        ["PUT", "/api/sop/execution/:exec_id", "실행 단계 업데이트 (checked_steps, memo)"],
        ["POST", "/api/sop/execution/:exec_id/complete", "실행 완료"],
        ["POST", "/api/sop/execution/:exec_id/broadcast", "상황전파 로그"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.7 보고서", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/reports", "보고서 목록"],
        ["GET", "/api/reports/:id", "보고서 상세"],
        ["POST", "/api/reports/generate", "이벤트 기반 자동생성 {event_id}"],
        ["PUT", "/api/reports/:id", "보고서 수정 (manager_comment 등)"],
        ["PATCH", "/api/reports/:id/status", "상태 변경 (DRAFT → SUBMITTED)"],
        ["DELETE", "/api/reports/:id", "삭제"],
        ["POST", "/api/reports/bulk-delete", "일괄 삭제 {report_ids[]}"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

add_heading(doc, "5.8 설정", level=2)
add_table(doc,
    ["Method", "Path", "용도"],
    [
        ["GET", "/api/settings", "설정 목록"],
        ["PUT", "/api/settings/:id", "설정 변경"],
        ["GET", "/api/settings/thresholds", "임계치 목록 (?equipment_id)"],
        ["PUT", "/api/settings/thresholds/:sensor_id", "임계치 수정"],
        ["GET", "/api/settings/sensor-meta", "센서 메타데이터"],
        ["PUT", "/api/settings/sensor-meta/:sensor_id", "센서 메타 수정"],
    ],
    col_widths=[2.0, 5.5, 9.5]
)

doc.add_page_break()

# ─────────────── 6. 시나리오 에뮬레이터 ───────────────
add_heading(doc, "6. 시나리오 에뮬레이터 요구사항", level=1)

add_heading(doc, "6.1 목적", level=2)
add_para(doc,
    "실데이터·실 외부기관 서버가 부재한 PoC 단계에서 8종 시나리오의 시계열 데이터를 재생하여 "
    "전체 서비스 흐름(센서 → 알람 → 이벤트 → 외부기관 결과 enrichment → SOP 추천 → 보고서 "
    "자동생성)을 시현한다. 본개발 단계에서는 \"시연/훈련 모드\"로 잔존하여 신규 사용자 교육·"
    "외부기관 연동 검증·UAT에 활용한다."
)

add_heading(doc, "6.2 시나리오 목록 (POC 8종)", level=2)
add_table(doc,
    ["ID", "시나리오 명", "Trigger", "주요 영향설비", "HAZOP"],
    [
        ["SC-01", "BOG 압축기 트립 → 저장탱크 압력 상승", "BOG-201", "TK-101, REL-701, VAL-601", "HZ-LH2-001"],
        ["SC-02", "2차 펌프 캐비테이션 → 기화기 영향", "PMP-301", "VAP-401, PIP-501", "HZ-LH2-002"],
        ["SC-03", "저장탱크 압력 상승 (BOG 급증)", "TK-101", "BOG-201, REL-701", "HZ-LH2-003"],
        ["SC-04", "벤트스택 밸브 Stuck Close", "VAL-601", "PMP-301, VAP-401, PIP-501", "HZ-LH2-004"],
        ["SC-05", "BOG 압축기 전단 저유량 (Surge)", "PIP-501", "BOG-201", "HZ-LH2-005"],
        ["SC-06", "기화기 후단 저온 취성 위험", "VAP-401", "PIP-501", "HZ-LH2-006"],
        ["SC-07", "정상 운전 (24h 기준 프로파일)", "TK-101", "—", "HZ-LH2-007"],
        ["SC-08", "로딩암 ESD (긴급차단)", "ARM-101", "SHP-001, VAL-601", "HZ-LH2-008"],
    ],
    col_widths=[1.5, 6.5, 2.0, 5.5, 2.0]
)

add_heading(doc, "6.3 Phase 진행 모델 (SC-01 기준)", level=2)
add_table(doc,
    ["Phase", "구간(sec)", "주요 동작"],
    [
        ["NORMAL", "0 ~ 179", "정상 센서값 송출, 전 설비 normal, 알람 없음"],
        ["SYMPTOM", "180 ~ 359", "초기 이상 센서값, trigger 설비 warning 컬러링, WARNING 알람 + 이벤트 팝업, KOGAS Mock + Safetia 이력 enrich"],
        ["FAULT", "360 ~ 599", "이상 확대, trigger 설비 critical/emergency, KGS 영향분석 → affected 설비 컬러링, SOP 추천 활성화, EventLog 생성"],
        ["SECONDARY_IMPACT", "600 ~ 779", "영향 설비 센서값 변화, affected 컬러링 확대, KETI 시뮬레이션 활성화"],
        ["RESPONSE", "780 ~ 900", "대응 완료, 점진적 normal 복귀, 이벤트 CLOSED, 보고서 초안 자동생성"],
    ],
    col_widths=[3.5, 3.0, 11.0]
)

add_heading(doc, "6.4 SSE 이벤트 스펙", level=2)
add_code_block(doc,
    "EmulatorEvent = {\n"
    "  type: 'CONNECTED' | 'SENSOR_UPDATE' | 'PHASE_CHANGE' | 'ALARM' |\n"
    "        'SYMPTOM_ENRICHMENT' | 'EVENT_CREATE' | 'EVENT_CLOSED' |\n"
    "        'REPORT_GENERATED' | 'SCENARIO_END' | 'SCENARIO_RESET' |\n"
    "        'EMULATOR_PAUSED' | 'EMULATOR_RESUMED',\n"
    "  timestamp: ISO8601,\n"
    "  phase: string,\n"
    "  elapsed_sec: number,\n"
    "  data: SensorDataBatch | PhaseInfo | AlarmInfo | EventInfo | EnrichmentInfo\n"
    "}"
)

doc.add_page_break()

# ─────────────── 7. 3D 시각화 요구사항 ───────────────
add_heading(doc, "7. 3D 시각화 요구사항", level=1)

add_heading(doc, "7.1 GLB 자산", level=2)
add_table(doc,
    ["파일", "용도", "크기", "Draco 압축", "정점 수", "호스팅"],
    [
        ["h2.glb", "테스트베드 기지 전체 (선박 포함)", "30 MB", "Required", "약 8.32M", "Cloudflare R2 런타임 로드"],
        ["secondary_pump.glb", "2차 펌프 상세 (M-ANO X-ray)", "210 KB", "미사용", "9,388", "앱 번들 (public/models)"],
    ],
    col_widths=[3.5, 5.5, 1.5, 2.0, 2.0, 2.5]
)

add_heading(doc, "7.2 컬러링 우선순위", level=2)
add_code_block(doc,
    "emergency  '#FF1744'   적색\n"
    "critical   '#FF5722'   주홍\n"
    "simTarget  '#E040FB'   보라  (시뮬레이션 대상)\n"
    "affected   '#FFEE58'   황색  (영향 설비)\n"
    "warning    '#FFA726'   주황\n"
    "normal     '#66BB6A'   녹색 (또는 기본 재질)\n"
    "\n"
    "우선순위 : emergency > critical > simTarget > affected > warning > normal"
)

add_heading(doc, "7.3 카메라 프리셋 (12종)", level=2)
add_table(doc,
    ["프리셋", "용도"],
    [
        ["cam_overview", "인수기지 전체 조감"],
        ["cam_berth_overview", "선석/하역부 전체 조망"],
        ["cam_ship_carrier_001", "운반선 정면"],
        ["cam_loading_arm_101", "로딩암 연결부"],
        ["cam_tank_101 / 102", "저장탱크 측면"],
        ["cam_bog_compressor_201", "BOG 압축기 정면"],
        ["cam_pump_301", "이송펌프 측면"],
        ["cam_vaporizer_401", "기화기 입출구"],
        ["cam_reliquefier_701", "재액화기 전면"],
        ["cam_valve_station_601 / 602", "벤트스택 밸브 정면"],
        ["cam_pipe_main_a", "메인 이송배관 전경"],
    ],
    col_widths=[5.5, 11.5]
)

add_heading(doc, "7.4 시각 이펙트 (9종)", level=2)
add_table(doc,
    ["이펙트", "활성화 조건", "구현"],
    [
        ["배관 유체 흐름", "항상 (모니터링/시뮬레이션)", "GLSL 셰이더 파티클 / UV 스크롤"],
        ["탱크 레벨/압력", "항상 (센서값 연동)", "ShaderMaterial uniforms"],
        ["설비 글로우", "WARNING 이상", "MeshBasicMaterial + AdditiveBlending"],
        ["히트맵 오버레이", "M-RSK, M-SIM 분석 후", "ShaderMaterial + transparent"],
        ["영향 전파 경로", "M-RSK KGS 결과 로드 후", "TubeGeometry + dash 애니메이션"],
        ["가스 확산 (M-SIM)", "수동 시뮬레이션 실행", "구체 + 분자 파티클 (deck.gl 컨셉)"],
        ["설비 POI 마커", "항상", "Three.js Sprite / HTML 오버레이"],
        ["환경 (sky/ground)", "항상", "EnvironmentScene"],
        ["X-ray 모드 (펌프)", "M-ANO PMP-301 선택 시", "secondary_pump.glb mesh 컬러링"],
    ],
    col_widths=[4.5, 5.5, 7.0]
)

add_heading(doc, "7.5 성능 가이드라인", level=2)
for t in [
    "파티클 수 : 배관당 최대 200개, 전체 1,000개 이내",
    "히트맵 텍스처 : 512×512 (런타임 생성)",
    "글로우 mesh : 원본 mesh 복제 후 scale 1.05x, 별도 렌더패스 미사용",
    "DPR (Device Pixel Ratio) : 모바일 1, 태블릿 1.5, 데스크톱 2",
    "안티앨리어싱 : 모바일 OFF, 그 외 ON",
    "카메라 far : 5,000 (h2.glb 좌표 범위 약 738 단위 대응)",
    "GPU 메모리 : Decimate 적용 시 약 60MB 목표 (현재 약 306MB)",
]:
    add_bullet(doc, t)

doc.add_page_break()

# ─────────────── 8. UI/UX 요구사항 ───────────────
add_heading(doc, "8. UI / UX 요구사항", level=1)

add_heading(doc, "8.1 공통 레이아웃", level=2)
add_code_block(doc,
    "[GNB] 로고 | M-MON M-ANO M-RSK M-SIM M-HIS M-SOP | 🔔알람(N) | ⚙설정 📋보고서 | 시나리오:SC-01 ▶️\n"
    "[API 상태바] 🟢KOGAS 정상 | 🟢KGS 정상 | 🟢KETI 정상 | 🟢세이프티아 정상\n"
    "[모드별 콘텐츠 영역]\n"
    "[하단 에뮬레이터 바] 진행: ████████░░ FAULT (6:23/15:00) | Speed: 10x | ⏸ ⏹"
)

add_heading(doc, "8.2 반응형 브레이크포인트", level=2)
add_table(doc,
    ["디바이스", "범위", "주요 변화"],
    [
        ["Mobile", "<= 767px", "GNB 햄버거 드로어, 하단 탭, 3D Canvas 50vh, KPI 수평 스크롤, SOP 팝업 → 전체화면"],
        ["Tablet", "768 ~ 1023px", "GNB 축약, 좌·우 패널 → 탭 전환, 2분할 레이아웃"],
        ["Desktop", ">= 1024px", "전체 3분할, 사이드바 드래그 리사이즈, 카메라 시점 저장/복원"],
    ],
    col_widths=[3.0, 3.0, 11.0]
)

add_heading(doc, "8.3 색상·아이콘 표준", level=2)
add_para(doc, "심각도 색상", bold=True)
add_table(doc,
    ["심각도", "색상", "Hex"],
    [
        ["EMERGENCY", "적색", "#FF1744"],
        ["CRITICAL", "주홍", "#FF5722"],
        ["WARNING", "주황", "#FFA726"],
        ["INFO / NORMAL", "녹색", "#66BB6A"],
    ],
    col_widths=[4.0, 4.0, 4.0]
)
add_para(doc, "설비 아이콘 (emoji 기반 — 본개발 시 디자인 시스템 아이콘셋으로 교체)", bold=True)
add_code_block(doc,
    "🚢 LH2_CARRIER     ⚓ LOADING_ARM    🏭 STORAGE_TANK\n"
    "💨 BOG_COMPRESSOR  🔧 TRANSFER_PUMP  🌡 VAPORIZER\n"
    "🔗 MAIN_PIPE       🔒 VALVE_STATION  ♻ RELIQUEFIER\n"
    "🌊 SEAWATER_PUMP"
)

doc.add_page_break()

# ─────────────── 9. 비기능 요구사항 ───────────────
add_heading(doc, "9. 비기능 요구사항", level=1)

add_heading(doc, "9.1 성능", level=2)
add_table(doc,
    ["항목", "PoC 수준", "본개발 목표"],
    [
        ["동시 접속", "5명 미만", "50명 이상 (운영기관·외부기관 합산)"],
        ["3D 초기 로드 시간", "3 ~ 10초 (h2.glb 30MB)", "5초 이내 (Decimate 적용 + CDN 캐시)"],
        ["센서 시계열 응답", "JSON 파일 직접 로드", "TSDB 쿼리 200ms 이내"],
        ["SSE 지연", "100 ms (틱 간격)", "100 ms 이내 유지"],
        ["KGS 분석 응답", "Mock 즉시", "5초 이내 (실 엔진)"],
        ["KETI 시뮬레이션", "Mock 즉시", "30초 이내 (실 엔진)"],
        ["보고서 생성", "1초 이내", "2초 이내"],
    ],
    col_widths=[5.0, 5.5, 6.5]
)

add_heading(doc, "9.2 가용성 / 신뢰성", level=2)
for t in [
    "외부기관 API 연결 실패 시: ApiStatusBar에 🔴 표시 + 마지막 응답 캐시로 fallback (TTL 정책 협의 필요)",
    "SSE 연결 단절 시: 자동 재연결(exponential backoff), 최대 5회",
    "DB 연결 실패 시: 헬스체크 503 응답, 자동 재시작 (Railway 정책 또는 운영기관 정책)",
    "데이터 품질: quality 필드(GOOD / ESTIMATED / MISSING)로 표기, MISSING 시 보간 또는 회색 처리",
]:
    add_bullet(doc, t)

add_heading(doc, "9.3 보안 (본개발 신규)", level=2)
for t in [
    "인증 : JWT 또는 OAuth2 (운영기관 IDP 연동)",
    "권한 : RBAC — 운영자 / 현장점검원 / 안전관리자 / 분석자 / 기관 검토자 / 관리자",
    "감사 로그 : 설정 변경, SOP 수정, 보고서 제출, 외부기관 호출 로그 영구 보관",
    "전송 보안 : 모든 API HTTPS, 외부기관 연동은 mTLS 권장",
    "외부기관 인증 : API Key 또는 Mutual TLS 인증서 (별첨 IFD §6 참조)",
    "민감정보 : 센서값 자체는 일반 데이터로 분류, 운영자 신원·결재 라인은 개인정보",
    "OWASP Top 10 대응 : SQL Injection (Prisma 파라미터 바인딩), XSS (React 기본 escape), CSRF (SameSite=Strict)",
]:
    add_bullet(doc, t)

add_heading(doc, "9.4 배포 / 운영", level=2)
add_table(doc,
    ["구성", "PoC", "본개발 옵션"],
    [
        ["Frontend", "Vercel (Next.js 14)", "(A) Vercel 유지 / (B) 운영기관 사내 호스팅"],
        ["Backend", "Railway (Node.js + Express)", "(A) Railway / (B) 운영기관 사내 K8s / (C) 클라우드 (AWS·NCP)"],
        ["DB", "Railway PostgreSQL", "(A) PostgreSQL 클러스터 / (B) AWS RDS / (C) NCP Cloud DB"],
        ["3D 자산", "Cloudflare R2", "(A) R2 / (B) 사내 CDN / (C) 운영기관 NAS"],
        ["TSDB (본개발 신규)", "—", "InfluxDB 또는 TimescaleDB"],
        ["로그/모니터링 (본개발 신규)", "—", "Grafana + Loki + Prometheus 또는 운영기관 표준 SIEM"],
    ],
    col_widths=[4.5, 5.0, 7.5]
)

add_heading(doc, "9.5 국제화 / 접근성", level=2)
for t in [
    "기본 언어 : 한국어 (모든 UI 텍스트 KR)",
    "본개발 확장 : 영문 동시 지원 (외부기관·해외 OEM 대비, i18n 키 분리)",
    "WCAG 2.1 AA 준수 (색맹/저시력 대응 — 색상 외 텍스트·아이콘 동시 표기)",
    "단위 표기 : SI 단위 일관 사용 (bar(g), °C, m³/h, mm/s, A, %)",
]:
    add_bullet(doc, t)

doc.add_page_break()

# ─────────────── 10. 본개발 확정 필요 항목 ───────────────
add_heading(doc, "10. 본개발 진입 전 확정 필요 항목", level=1)
add_para(doc,
    "본 절은 PoC를 검토한 각 기관·운영기관이 본개발 착수 전까지 의사결정을 내려야 할 항목을 "
    "정리한다. 「기관별 데이터 인터페이스 정의서(IFD)」와 함께 검토되어야 한다."
)

add_heading(doc, "10.1 모니터링 데이터 정의 — 운영기관 + 4개 기관 공통", level=2)
add_table(doc,
    ["#", "항목", "확정 주체", "PoC 가정", "확정 필요사항"],
    [
        ["1", "전수 설비 ID 체계", "운영기관 + 전 기관", "12종 핵심설비 ID (TK-101 등)", "실 P&ID 기반 전수 설비 ID + tag-naming 표준"],
        ["2", "전수 센서 ID 체계", "운영기관", "36종 (PRESSURE/TEMP/FLOW/VIB/CURRENT/LEVEL)", "실제 센서 전수 + KKS·ISA 표준 적용 여부"],
        ["3", "센서 샘플링 주기", "운영기관", "5초 (기본)", "센서 유형별 차등 (진동 100ms, 압력 1초 등)"],
        ["4", "임계치 운영 기준", "운영기관 + 안전관리자", "PoC seed 값 (정상값 ±15/30%)", "운영기준치 + 계절·운전조건별 동적 임계치"],
        ["5", "데이터 보존 기간", "운영기관", "PoC 비적용", "센서 raw / 1분 평균 / 1시간 평균 / 일별 / 월별 보존 정책"],
        ["6", "이벤트 발생 정의", "전 기관", "센서 임계치 초과 또는 KOGAS push", "확정 알고리즘 (단일 센서 / 다중 센서 AND/OR / 시간창 등)"],
        ["7", "데이터 품질 정책", "운영기관", "GOOD / ESTIMATED / MISSING", "MISSING 시 처리 (보간 / 알람 / SOP 트리거)"],
    ],
    col_widths=[0.7, 4.5, 3.0, 4.0, 4.8]
)

add_heading(doc, "10.2 외부기관 인터페이스 — IFD 동기화", level=2)
add_table(doc,
    ["기관", "확정 필요사항", "관련 IFD 절"],
    [
        ["KOGAS", "이상감지 push/pull 방식, 진단 응답시간, sensor_evidence 형식(배열 string vs 객체), fault_code 체계, 진단 신뢰도 기준값", "IFD §4"],
        ["KGS", "위험예측 입력 파라미터 표준, impact_score 산출 기준, predicted_after_sec 정밀도, 영향설비 그래프 깊이, recommended_action 텍스트 표준", "IFD §5"],
        ["KETI", "시뮬레이션 입력 파라미터 표준, 대응안 개수(A/B 또는 N개), simulation_summary 문장 길이 제한, 동적 저작도구 연동 여부", "IFD §6"],
        ["세이프티아", "이력 조회 페이지네이션·필터·기간 검색, 사고대응 라이브러리 연동, 이력 영구 동기화 vs 매 호출", "IFD §7"],
    ],
    col_widths=[2.5, 11.0, 3.5]
)

add_heading(doc, "10.3 SOP 정책", level=2)
for t in [
    "운영기관 보유 SOP 전수 등록 절차 (포맷 변환·번역·승인)",
    "SOP 버전 관리 정책 (Major/Minor 변경 기준, 이전 버전 보존)",
    "결재 라인 적용 여부 (작성 → 검토 → 승인 단계)",
    "현장 점검 시 첨부파일·사진·전자서명 요건",
    "비상 SOP 자동팝업 정책 (자동 vs 수동 확인, 음성 알림 등)",
]:
    add_bullet(doc, t)

add_heading(doc, "10.4 보고서 정책", level=2)
for t in [
    "보고서 템플릿 확장 (RPT-TPL-001 외에 일일·주간·월간·사고조사·점검·정비 등)",
    "PDF 출력 양식 (운영기관 표준 양식 적용)",
    "결재 라인 적용 여부, 외부 기관 회람 여부",
    "감사·법규 대응 자료 자동 추출 기능",
]:
    add_bullet(doc, t)

add_heading(doc, "10.5 운영 모드 / 시연 모드 분리", level=2)
for t in [
    "본개발 단계에서는 \"실 운영 모드\"(실데이터 + 실 외부기관) 와 \"시연/훈련 모드\"(에뮬레이터)를 명확히 분리",
    "권한 분리: 시연 모드는 \"기관 검토자\"·\"신규 사용자\" 전용, 실데이터에 영향 미치지 않음",
    "이력·이벤트·보고서는 모드별로 분리 저장 (가짜 데이터로 인한 운영 통계 오염 방지)",
]:
    add_bullet(doc, t)

doc.add_page_break()

# ─────────────── 11. 본개발 단계 로드맵 ───────────────
add_heading(doc, "11. 본개발 단계 로드맵 (제안)", level=1)
add_table(doc,
    ["단계", "범위", "주요 산출물", "기관 협업"],
    [
        ["Phase 0\n(준비)", "본 요구사항 정의서 + IFD 검토 회의 (4개 기관)", "확정안 v1.0", "전 기관"],
        ["Phase 1\n(인프라)", "운영기관 인프라 결정, 보안정책 적용, CI/CD", "배포 환경 구축", "운영기관"],
        ["Phase 2\n(데이터 표준)", "전수 설비/센서 ID 표준, 임계치, 마스터 데이터 적재", "마스터 DB + ETL", "운영기관"],
        ["Phase 3\n(외부기관 연동)", "KOGAS/KGS/KETI/세이프티아 실 API 연동, IFD 단계별 검증", "연동 테스트 보고서", "전 기관"],
        ["Phase 4\n(통합 UI)", "8개 모드 본개발, RBAC 적용, 결재 라인", "Frontend v1.0", "유엔이 + 운영기관"],
        ["Phase 5\n(UAT)", "운영기관 UAT, 외부기관 합동 시연", "UAT 보고서", "전 기관"],
        ["Phase 6\n(상용 운영)", "이행, 모니터링, 운영 매뉴얼", "운영 가이드", "운영기관"],
    ],
    col_widths=[2.5, 5.0, 4.0, 5.5]
)

doc.add_page_break()

# ─────────────── 12. 부록 ───────────────
add_heading(doc, "12. 부록", level=1)

add_heading(doc, "12.1 PoC 시드 데이터 파일 목록 (28개)", level=2)
add_code_block(doc,
    "seed_master_zone.json                  (8 zones)\n"
    "seed_master_equipment.json             (12 equipment)\n"
    "seed_master_sensor_type.json           (6 types)\n"
    "seed_master_sensor.json                (36 sensors)\n"
    "seed_equipment_sensor_map.json         (38 mappings)\n"
    "seed_sensor_thresholds.json            (37 thresholds)\n"
    "seed_mock_scenarios.json               (8 scenarios)\n"
    "seed_hazop_lh2.json                    (8 HAZOP entries)\n"
    "seed_event_log.json                    (8 events)\n"
    "seed_mock_kogas_results.json           (8 KOGAS results)\n"
    "seed_mock_kgs_results.json             (25 KGS results)\n"
    "seed_mock_keti_results.json            (8 KETI results)\n"
    "seed_mock_safetia_history.json         (8 Safetia histories)\n"
    "seed_sop_catalog.json                  (9 SOPs)\n"
    "seed_sop_equipment_map.json            (17 mappings)\n"
    "seed_sop_execution_samples.json        (8 executions)\n"
    "seed_report_templates.json             (1 template)\n"
    "seed_report_samples.json               (7 reports)\n"
    "seed_settings_metadata.json            (5 settings)\n"
    "seed_pump_mesh_coloring.json           (SC-02 pump mesh coloring rules)\n"
    "seed_process_stages.json               (4 process stages)\n"
    "seed_sensor_timeseries_SC-01 ~ 08.json (8 timeseries files)\n"
    "seed_manifest.json"
)

add_heading(doc, "12.2 코드 체계 (Enum)", level=2)
add_table(doc,
    ["분류", "허용값"],
    [
        ["EquipmentType", "LH2_CARRIER, LOADING_ARM, STORAGE_TANK, BOG_COMPRESSOR, TRANSFER_PUMP, VAPORIZER, MAIN_PIPE, VALVE_STATION, RELIQUEFIER, SEAWATER_PUMP"],
        ["SensorType", "PRESSURE, TEMPERATURE, FLOW, VIBRATION, CURRENT, LEVEL"],
        ["Severity", "INFO, WARNING, CRITICAL, EMERGENCY"],
        ["VisualState", "normal, warning, affected, critical, emergency, simTarget"],
        ["EventStatus", "OPEN → PROCESSING → CLOSED"],
        ["SopCategory", "EMERGENCY, SAFETY, ROUTINE, INSPECTION, EVENT_RESPONSE, MAINTENANCE (PoC)"],
        ["SopStepType", "TEXT, CHECK, DECISION (분기)"],
        ["ReportStatus", "DRAFT, SUBMITTED"],
        ["SensorQuality", "GOOD, ESTIMATED, MISSING"],
        ["SensorLabel", "NORMAL, WARNING, ANOMALY"],
        ["Phase", "NORMAL, SYMPTOM, FAULT, SECONDARY_IMPACT, RESPONSE"],
    ],
    col_widths=[3.5, 13.5]
)

add_heading(doc, "12.3 변경 이력", level=2)
add_table(doc,
    ["버전", "일자", "주요 변경", "작성자"],
    [
        ["v1.0", "2026-05-07", "PoC 역설계 → 본개발 기준안 최초 작성", "유엔이"],
    ],
    col_widths=[2.0, 3.0, 9.0, 3.0]
)

add_heading(doc, "12.4 검토자 서명란", level=2)
add_table(doc,
    ["기관", "검토자(직책/성명)", "검토일", "의견 / 서명"],
    [
        ["KOGAS", "", "", ""],
        ["KGS", "", "", ""],
        ["KETI", "", "", ""],
        ["세이프티아", "", "", ""],
        ["운영기관", "", "", ""],
        ["유엔이(작성)", "", "", ""],
    ],
    col_widths=[3.0, 4.5, 3.0, 6.5]
)

# 저장
out_path = "/home/user/UNE/docs/REQ_SPEC_LH2_DigitalTwin_v1.0.docx"
doc.save(out_path)
print(f"[OK] saved: {out_path}")
